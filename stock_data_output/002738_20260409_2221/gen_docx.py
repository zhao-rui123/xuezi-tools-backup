#!/usr/bin/env python3
"""中矿资源股票分析报告 -> Word文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

output_dir = "/Users/zhaoruicn/.openclaw/workspace/stock_data_output/002738_20260409_2221"
doc_path = os.path.join(output_dir, "中矿资源_002738_分析报告.docx")

doc = Document()

# 标题
title = doc.add_heading('中矿资源 (002738.SZ) 股票分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('分析时间：2026-04-09 | 当前股价：78.0元（+0.31%）| 市值：~562亿')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 一、技术面分析
doc.add_heading('一、技术面分析', level=1)

# 添加图片
img_path = os.path.join(output_dir, "kline_em.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图1：日K线图 — 近期走势').alignment = WD_ALIGN_PARAGRAPH.CENTER

intraday_path = os.path.join(output_dir, "kline_intraday.png")
if os.path.exists(intraday_path):
    doc.add_picture(intraday_path, width=Inches(6))
    doc.add_paragraph('图2：分时图 — 今日走势').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_heading('1.1 走势概览', level=2)
doc.add_paragraph('股价自2025年中低点27元附近持续上涨，2026年1月创出100.86元高点后回落整固，目前在75-79元区间震荡，4月9日上涨0.31%收于78元，逼近前期高点。')

doc.add_heading('1.2 均线系统', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['均线', '数值', '状态']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('MA5', '75.24', '向上'),
    ('MA20', '72.09', '向上'),
    ('MA60', '81.0', '向下'),
    ('MA5/MA20', '交叉缠绕', '中性'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph('均线系统显示：短期均线向上，中期均线向下，整体震荡格局。')

doc.add_heading('1.3 支撑与压力', level=2)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
for i, h in enumerate(['类型', '位置', '意义']):
    table2.rows[0].cells[i].text = h
data2 = [
    ('支撑1', '75元（MA5）', '近期低点，需守住'),
    ('支撑2', '72元（MA20）', '中期支撑'),
    ('压力1', '78.69元', '52周高点，需放量突破'),
    ('压力2', '81元（MA60）', '中期均线压力'),
]
for i, row in enumerate(data2):
    for j, val in enumerate(row):
        table2.rows[i+1].cells[j].text = val

doc.add_heading('1.4 今日分时解读', level=2)
doc.add_paragraph('• 开盘：77.76元（前一日收盘）')
doc.add_paragraph('• 最高：78.69元（涨1.2%，接近52周高点）')
doc.add_paragraph('• 最低：77.76元')
doc.add_paragraph('• 收盘：78.0元（+0.31%）')
doc.add_paragraph('• 成交量：0.58倍均量，量能萎缩')

doc.add_heading('1.5 关键技术信号', level=2)
table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Table Grid'
for i, h in enumerate(['指标', '数值', '信号']):
    table3.rows[0].cells[i].text = h
data3 = [
    ('RSI', '55.8', '中性（50上方偏多）'),
    ('MACD', '动能增强', 'DIF向上，红柱放大'),
    ('MA状态', '交叉缠绕', '震荡'),
    ('量能', '0.58倍', '缩量，需补量确认'),
]
for i, row in enumerate(data3):
    for j, val in enumerate(row):
        table3.rows[i+1].cells[j].text = val

p = doc.add_paragraph()
run = p.add_run('技术面评级：🟡 中性偏强（MACD向好，但量能不足）')
run.bold = True

doc.add_paragraph()

# 二、基本面分析
doc.add_heading('二、基本面分析', level=1)

doc.add_heading('2.1 核心财务数据', level=2)
table4 = doc.add_table(rows=5, cols=3)
table4.style = 'Table Grid'
for i, h in enumerate(['指标', '数值', '备注']):
    table4.rows[0].cells[i].text = h
data4 = [
    ('市盈率 (PE)', '122.97', '较高'),
    ('市净率 (PB)', '4.62', '中性'),
    ('52周高点', '78.69元', '当前接近'),
    ('52周低点', '27.01元', '已大涨190%'),
]
for i, row in enumerate(data4):
    for j, val in enumerate(row):
        table4.rows[i+1].cells[j].text = val

doc.add_heading('2.2 2025年业绩', level=2)
table5 = doc.add_table(rows=3, cols=5)
table5.style = 'Table Grid'
for i, h in enumerate(['报告期', '营收', '同比', '净利', '同比']):
    table5.rows[0].cells[i].text = h
data5 = [
    ('2025全年', '48.18亿', '+34.99%', '2.04亿', '-62.58%'),
    ('2025Q3', '15.51亿', '+35.19%', '1.15亿', '+58.18%'),
]
for i, row in enumerate(data5):
    for j, val in enumerate(row):
        table5.rows[i+1].cells[j].text = val

p = doc.add_paragraph()
run = p.add_run('关键点：')
run.bold = True
doc.add_paragraph('营收增长但净利下滑，主要因锂价下跌压制盈利能力。')

doc.add_heading('2.3 业务结构', level=2)
doc.add_paragraph('中矿资源主营：')
doc.add_paragraph('• 锂电新材料：碳酸锂、氢氧化锂')
doc.add_paragraph('• 稀有金属：铯、铷（全球垄断优势）')
doc.add_paragraph('• 地勘技术服务：传统业务')

doc.add_heading('2.4 研报观点', level=2)
doc.add_paragraph('最新券商评级（2026年4月）：')
doc.add_paragraph('• 东吴证券：买入（机器人+航空航天打开增量空间）')
doc.add_paragraph('• 华源证券：买入（毛利率持续改善，在手订单超千亿）')
doc.add_paragraph('• 国信证券：增持（资源端和材料端持续突破）')

p = doc.add_paragraph()
run = p.add_run('基本面评级：🟡 中性（营收增长、券商看好，但PE偏高、净利下滑）')
run.bold = True

doc.add_paragraph()

# 三、资金流向
doc.add_heading('三、资金流向', level=1)
table6 = doc.add_table(rows=3, cols=3)
table6.style = 'Table Grid'
for i, h in enumerate(['日期', '主力净流入', '散户净流入']):
    table6.rows[0].cells[i].text = h
table6.rows[1].cells[0].text = '4月9日'
table6.rows[1].cells[1].text = '-4191万'
table6.rows[1].cells[2].text = '+1400万'

doc.add_paragraph()
doc.add_paragraph('融资余额：31.77亿元（4月8日），环比+3.30%，融资客在加仓。')

p = doc.add_paragraph()
run = p.add_run('资金面评级：🔴 偏弱（主力净流出，但融资余额上升）')
run.bold = True

doc.add_paragraph()

# 四、综合判断
doc.add_heading('四、综合判断', level=1)

doc.add_heading('4.1 多空对比', level=2)
table7 = doc.add_table(rows=5, cols=2)
table7.style = 'Table Grid'
table7.rows[0].cells[0].text = '🟢 做多逻辑'
table7.rows[0].cells[1].text = '🔴 做空逻辑'
table7.rows[1].cells[0].text = '锂价修复，Q3净利大增58%'
table7.rows[1].cells[1].text = 'PE高达123倍，估值极贵'
table7.rows[2].cells[0].text = '券商普遍买入评级'
table7.rows[2].cells[1].text = '2025全年净利下滑62%'
table7.rows[3].cells[0].text = '融资余额持续增长'
table7.rows[3].cells[1].text = '主力资金连续净流出'
table7.rows[4].cells[0].text = '铯铷全球垄断，护城河深'
table7.rows[4].cells[1].text = '成交量萎缩，上涨动力不足'

doc.add_heading('4.2 核心结论', level=2)
p = doc.add_paragraph()
run = p.add_run('总评级：🟡 观望')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph('理由：')
doc.add_paragraph('1. 技术面MACD向好，但量能不足，突破需要补量')
doc.add_paragraph('2. 基本面营收增长但净利下滑，PE极高（123倍）')
doc.add_paragraph('3. 券商看好但目标价与现价差距大（历史目标价45-48元，当前78元已超预期）')
doc.add_paragraph('4. 主力资金连续净流出，筹码分散')

doc.add_heading('4.3 操作建议', level=2)
doc.add_paragraph('• 已持仓者：可持有，但需设78.69元止损，跌破75考虑减仓')
doc.add_paragraph('• 观望者：暂不追高，等回踩72-75元支撑再考虑')
doc.add_paragraph('• 短线者：78-79元区间高抛低吸，破75止损')

doc.add_paragraph()

# 五、风险提示
doc.add_heading('五、风险提示', level=1)
doc.add_paragraph('1. 锂价波动风险：公司盈利与锂价高度相关')
doc.add_paragraph('2. 估值风险：PE 123倍，股价已充分反映乐观预期')
doc.add_paragraph('3. 解禁风险：需关注股东减持动态')
doc.add_paragraph('4. 市场情绪：资源品周期波动大')

doc.add_paragraph()
doc.add_paragraph()

# 页脚
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('报告由雪子助手生成 | 数据来源：东方财富、券商研报')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(doc_path)
print(f"✅ Word文档已生成: {doc_path}")
