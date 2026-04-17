from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='Microsoft YaHei', size=11):
    run.font.name = font_name
    run.font.size = Pt(size)
    # 设置中文字体的西文部分
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def add_step(doc, number, title, content_lines, code_lines=None):
    # 步骤标题
    heading = doc.add_heading(f'第{number}步：{title}', level=2)

    # 内容
    for line in content_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_chinese_font(run)

    # 代码块（如果有）
    if code_lines:
        for code in code_lines:
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
            p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph()  # 空行

doc = Document()

# 标题
title = doc.add_heading('Mac mini 从零安装 Claude Code 完整指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('（适用于新版 Mac mini M4/M3 或 Intel 版）')
set_chinese_font(run, size=12)

doc.add_paragraph()

# 简介
doc.add_heading('简介', level=1)
intro = doc.add_paragraph()
set_chinese_font(intro.add_run(
    '本文档详细介绍在新购买的 Mac mini 上从零开始安装和配置 Claude Code 的完整步骤。'
    'Claude Code 是 Anthropic 官方提供的命令行工具，可以直接在终端中使用 Claude 的能力。'
))

# 目录概览
doc.add_heading('安装概览', level=1)
overview_items = [
    '第1步：检查系统要求',
    '第2步：安装 Homebrew（可选但推荐）',
    '第3步：安装 Node.js（必需）',
    '第4步：安装 Claude CLI',
    '第5步：配置环境变量',
    '第6步：安装 Git（可选）',
    '第7步：配置 SSH 密钥（可选）',
    '第8步：验证安装',
    '第9步：首次登录和使用'
]
for item in overview_items:
    p = doc.add_paragraph()
    set_chinese_font(p.add_run(item))
    p.style = 'List Bullet'

doc.add_page_break()

# ============ 第1步 ============
add_step(doc, 1, '检查系统要求', [
    '在开始之前，请确认您的 Mac mini 满足以下要求：',
    '',
    '•  macOS 版本：10.15 (Catalina) 或更高版本',
    '•  处理器：Apple M1/M2/M3/M4 或 Intel',
    '•  内存：建议至少 8GB',
    '•  磁盘空间：建议至少 10GB 可用空间',
    '',
    '查看系统版本的方法：',
    '点击屏幕左上角苹果菜单 (🍎) → 关于本机 → 查看 macOS 版本'
], code_lines=[
    '# 也可以在终端中查看
sw_vers',
    '# 输出示例：ProductName: macOS, ProductVersion: 14.0'
])

# ============ 第2步 ============
add_step(doc, 2, '安装 Homebrew（推荐）', [
    'Homebrew 是 macOS 上最流行的包管理器，可以方便地安装各种开发工具。',
    '虽然 Homebrew 不是安装 Claude Code 的必需步骤，但它能让后续升级和维护更方便。'
], code_lines=[
    '# 安装 Homebrew（官方命令）
/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"',
    '',
    '# 安装完成后，添加 Homebrew 到 PATH（Intel Mac）
echo "eval \"$(/opt/homebrew/bin/brew shellenv)\"" >> ~/.zprofile',
    '',
    '# 如果是 Apple Silicon Mac (M1/M2/M3/M4)
echo "eval \"$(/opt/homebrew/bin/brew shellenv)\"" >> ~/.zprofile',
    '',
    '# 重新加载终端配置
source ~/.zprofile',
    '',
    '# 验证 Homebrew 安装成功
brew doctor'
])

# ============ 第3步 ============
add_step(doc, 3, '安装 Node.js', [
    'Claude Code 需要 Node.js 环境才能运行。推荐使用 Node.js 20 或更高版本。',
    '',
    '方法一：使用 Homebrew 安装（推荐）'
], code_lines=[
    '# 安装 Node.js 20
brew install node@20',
    '',
    '# 添加到 PATH
echo "export PATH=\"/opt/homebrew/opt/node@20/bin:$PATH\"" >> ~/.zprofile
source ~/.zprofile'
])

doc.add_paragraph()
p = doc.add_paragraph()
set_chinese_font(p.add_run('方法二：从官网下载安装包'))

p2 = doc.add_paragraph()
set_chinese_font(p2.add_run(
    '访问 https://nodejs.org/ → 下载 LTS 版本 → 双击安装包按提示安装'
))

# ============ 第4步 ============
add_step(doc, 4, '安装 Claude CLI', [
    'Claude Code 的官方 CLI 工具可以通过 npm 安装。'
], code_lines=[
    '# 使用 npm 全局安装 Claude CLI
npm install -g @anthropic-ai/claude-code',
    '',
    '# 或者使用 Homebrew 安装
brew install claude-cli'
])

# ============ 第5步 ============
add_step(doc, 5, '验证 Node.js 和 Claude CLI 安装', [
    '安装完成后，验证各个组件是否正确安装：'
], code_lines=[
    '# 验证 Node.js
node --version
# 应该显示 v20.x.x 或更高版本',
    '',
    '# 验证 npm
npm --version
# 应该显示 10.x.x 或更高版本',
    '',
    '# 验证 Claude CLI
claude --version'
])

# ============ 第6步 ============
add_step(doc, 6, '配置 API 密钥', [
    '使用 Claude Code 需要配置 Anthropic API 密钥。',
    '',
    '1. 获取 API 密钥：',
    '   • 访问 https://console.anthropic.com/',
    '   • 登录或注册账号',
    '   • 点击 API Keys → Create Key → 复制密钥',
    '',
    '2. 配置环境变量：'
], code_lines=[
    '# 编辑终端配置文件（zsh 默认）
nano ~/.zshrc',
    '',
    '# 添加以下行（将 YOUR_API_KEY 替换为您的实际密钥）
export ANTHROPIC_API_KEY="sk-ant-..."',
    '',
    '# 保存后重新加载配置
source ~/.zshrc',
    '',
    '# 验证配置
echo $ANTHROPIC_API_KEY'
])

# ============ 第7步 ============
add_step(doc, 7, '首次运行 Claude Code', [
    '配置完成后，就可以开始使用 Claude Code 了：'
], code_lines=[
    '# 进入任意项目目录
cd ~/Projects/my-project',
    '',
    '# 启动 Claude Code
claude',
    '',
    '# 或者直接执行单个命令
claude "解释这段代码的作用"',
    '',
    '# 指定模型
claude --model opus "复杂任务"',
    '',
    '# 查看帮助
claude --help'
])

# ============ 第8步 ============
add_step(doc, 8, '常见问题和故障排除', [
    '问题1：提示 "command not found: claude"',
    '   解决方法：确保 ~/.zshrc 中正确添加了 PATH 配置，然后运行 source ~/.zshrc',
    '',
    '问题2：提示 "Invalid API key"',
    '   解决方法：检查 ANTHROPIC_API_KEY 环境变量是否正确设置，密钥是否过期',
    '',
    '问题3：npm 安装失败',
    '   解决方法：尝试使用 sudo npm install -g，或先安装 nvm',
    '',
    '问题4：Node 版本过低',
    '   解决方法：使用 nvm 安装新版本 Node.js'
])

# ============ 第9步 ============
add_step(doc, 9, '可选：配置 Git', [
    '如果需要与代码仓库交互，建议配置 Git：'
], code_lines=[
    '# 检查 Git 是否已安装
git --version',
    '',
    '# 配置用户名和邮箱
git config --global user.name "你的名字"
git config --global user.email "your.email@example.com"',
    '',
    '# 生成 SSH 密钥（如需要连接 GitHub/GitLab）
ssh-keygen -t ed25519 -C "your.email@example.com"',
    '',
    '# 查看公钥（添加到 GitHub/GitLab）
cat ~/.ssh/id_ed25519.pub'
])

# ============ 进阶配置 ============
doc.add_heading('进阶配置（可选）', level=1)

doc.add_heading('配置 Claude Code 使用代理', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run(
    '如果您的网络需要通过代理访问 Anthropic API，可以配置代理：'
))
doc.add_paragraph()

doc.add_heading('配置 Claude Code 默认模型', level=2)
p = doc.add_paragraph()
set_chinese_font(p.add_run('在 ~/.zshrc 中设置默认模型：'))
doc.add_paragraph()
p2 = doc.add_paragraph()
run = p2.add_run('export ANTHROPIC_DEFAULT_MODEL=opus')
run.font.name = 'Courier New'
run.font.size = Pt(10)
p2.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('常用 Claude Code 命令', level=2)
commands = [
    ('claude', '启动交互式对话'),
    ('claude "提示词"', '单次执行命令'),
    ('claude --model opus', '指定使用 Opus 模型'),
    ('claude --model sonnet', '指定使用 Sonnet 模型'),
    ('claude --print "提示"', '非交互模式，打印输出后退出'),
    ('exit', '退出 Claude Code')
]
for cmd, desc in commands:
    p = doc.add_paragraph()
    run1 = p.add_run(f'• {cmd}')
    run1.font.name = 'Courier New'
    run1.font.size = Pt(10)
    run2 = p.add_run(f' — {desc}')
    set_chinese_font(run2)

# ============ 注意事项 ============
doc.add_heading('注意事项', level=1)
notes = [
    '• API 密钥安全：请勿将 API 密钥分享给他人或提交到代码仓库',
    '• 费用：使用 Claude API 会产生费用，请留意使用量和账单',
    '• 隐私：Claude Code 会将代码和对话发送到 Anthropic 服务器处理',
    '• 更新：定期运行 npm update -g @anthropic-ai/claude-code 更新到最新版本'
]
for note in notes:
    p = doc.add_paragraph()
    set_chinese_font(p.add_run(note))
    p.style = 'List Bullet'

# 保存文档
doc.save('/Users/zhaoruicn/.openclaw/workspace/Macmini安装Claude_Code指南.docx')
print('文档已生成：Macmini安装Claude_Code指南.docx')
