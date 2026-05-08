"""Design-institute level DOCX report generator for energy solution analysis.

Produces professional Word documents with cover page, KPI dashboard, charts,
tables, and callout boxes — suitable for client delivery.
"""

from __future__ import annotations

import math
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from docx.table import _Cell

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import matplotlib.ticker as mticker
import numpy as np

# ═══════════════════════════════════════════════════════════════════
# Font setup
# ═══════════════════════════════════════════════════════════════════

_CHINESE_FONTS = [
    "Microsoft YaHei", "SimHei", "SimSun", "Source Han Sans SC",
    "Noto Sans CJK SC", "PingFang SC", "STHeiti", "DengXian",
]

def _find_chinese_font() -> str | None:
    for name in _CHINESE_FONTS:
        try:
            prop = fm.FontProperties(family=name)
            if prop.get_name():
                return name
        except Exception:
            continue
    return None

_CN_FONT = _find_chinese_font()
if _CN_FONT:
    plt.rcParams["font.family"] = _CN_FONT
    plt.rcParams["axes.unicode_minus"] = False

# ═══════════════════════════════════════════════════════════════════
# Color system — architecture-firm navy palette
# ═══════════════════════════════════════════════════════════════════

class C:
    NAVY      = RGBColor(0x0B, 0x1F, 0x3F)  # deepest
    DARK_BLUE = RGBColor(0x14, 0x2D, 0x5E)
    MID_BLUE  = RGBColor(0x1F, 0x47, 0x8A)
    ACCENT    = RGBColor(0x2E, 0x6B, 0xC4)
    LIGHT_BLUE= RGBColor(0x4A, 0x90, 0xD9)
    PALE_BLUE = RGBColor(0x8E, 0xBE, 0xEB)
    ICE       = RGBColor(0xD9, 0xE8, 0xF7)
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

    GREEN     = RGBColor(0x1B, 0x8C, 0x4A)
    LIGHT_GREEN=RGBColor(0xDF, 0xF0, 0xE3)
    RED       = RGBColor(0xC0, 0x39, 0x2B)
    LIGHT_RED = RGBColor(0xF9, 0xE3, 0xE0)
    ORANGE    = RGBColor(0xD6, 0x84, 0x0B)
    LIGHT_ORANGE=RGBColor(0xFE, 0xF2, 0xD6)

    GRAY_DARK = RGBColor(0x4A, 0x4A, 0x4A)
    GRAY      = RGBColor(0x88, 0x88, 0x88)
    GRAY_LIGHT= RGBColor(0xBB, 0xBB, 0xBB)
    BG_GRAY   = RGBColor(0xF5, 0xF6, 0xFA)

    BLACK     = RGBColor(0x1A, 0x1A, 0x1A)

PLOT_COLORS = ["#1F478A", "#1B8C4A", "#C0392B", "#D6840B", "#7D3C98",
               "#148F77", "#B9770E", "#2471A3", "#27AE60", "#8E44AD"]

PLOT_BLUE = "#1F478A"
PLOT_GREEN = "#1B8C4A"
PLOT_RED = "#C0392B"
PLOT_ORANGE = "#D6840B"

HEX_NAVY = "0B1F3F"
HEX_DARK_BLUE = "142D5E"
HEX_MID_BLUE = "1F478A"
HEX_ICE = "D9E8F7"
HEX_BG = "F5F6FA"
HEX_WHITE = "FFFFFF"
HEX_LIGHT_GREEN = "DFF0E3"
HEX_LIGHT_RED = "F9E3E0"
HEX_LIGHT_ORANGE = "FEF2D6"


# ═══════════════════════════════════════════════════════════════════
# Core helpers
# ═══════════════════════════════════════════════════════════════════

def _fmt(v: Any, decimals: int = 2) -> str:
    if v is None: return "—"
    if isinstance(v, bool): return "是" if v else "否"
    try:
        fv = float(v)
        if abs(fv) >= 1e8: return f"{fv:,.2f}"
        if abs(fv) >= 1e4: return f"{fv:,.{decimals}f}"
        if decimals == 0: return f"{fv:,.0f}"
        return f"{fv:,.{decimals}f}"
    except (TypeError, ValueError):
        return str(v)

def _pct(v: Any) -> str:
    if v is None: return "—"
    if isinstance(v, bool): return "是" if v else "否"
    try:
        fv = float(v)
        return f"{fv * 100:.2f}%" if abs(fv) < 1 else f"{fv:.2f}%"
    except (TypeError, ValueError):
        return str(v)

def _make_run(paragraph, text: str, bold: bool = False,
              size: int = 10, color: RGBColor = None,
              font_name: str = None) -> Any:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or C.BLACK
    run.font.name = font_name or (_CN_FONT or "Calibri")
    rPr = run._element.get_or_add_rPr()
    rPr.set(qn("w:eastAsia"), _CN_FONT or "")
    return run


# ── Table helpers ─────────────────────────────────────────────────

def _cell_shading(cell: _Cell, color: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def _cell_text(cell: _Cell, text: str, bold: bool = False,
               size: int = 9, color: RGBColor = None,
               align: str = "left", font_name: str = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _make_run(p, text, bold=bold, size=size, color=color or C.BLACK, font_name=font_name)

def _cell_v_align(cell: _Cell, align: str = "center") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def _add_table(doc: Document, headers: list[str],
               rows: list[list[str]], col_widths: list[float] | None = None,
               header_color: str = HEX_MID_BLUE,
               first_col_bold: bool = False) -> None:
    """Professional table with navy header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Striped flag
    for r_idx in range(1, len(table.rows)):
        if r_idx % 2 == 1:
            for cell in table.rows[r_idx].cells:
                _cell_shading(cell, HEX_BG)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _cell_shading(cell, header_color)
        _cell_text(cell, h, bold=True, size=9, color=C.WHITE, align="center")
        _cell_v_align(cell, "center")

    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            is_first = c_idx == 0
            _cell_text(cell, str(val), bold=(first_col_bold and is_first),
                       size=9, align="center" if c_idx > 0 else "left")
            _cell_v_align(cell, "center")

    if col_widths:
        for row_obj in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row_obj.cells):
                    row_obj.cells[i].width = Cm(w)

    doc.add_paragraph()


def _add_kpi_cards(doc: Document, cards: list[tuple[str, str, str, RGBColor]]) -> None:
    """Horizontal KPI cards in colored boxes. Each card: (label, value, note, color)."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (label, value, note, accent) in enumerate(cards):
        cell = table.rows[0].cells[i]
        hex_color = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
        _cell_shading(cell, hex_color)
        _cell_v_align(cell, "center")

        # Clear and add styled paragraphs
        cell.text = ""
        # Label
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(4)
        _make_run(p1, label, bold=False, size=8, color=C.WHITE)

        # Value
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _make_run(p2, value, bold=True, size=16, color=C.WHITE)

        # Note
        if note:
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _make_run(p3, note, bold=False, size=7, color=C.WHITE)

    doc.add_paragraph()


def _callout(doc: Document, text: str, color: RGBColor = C.MID_BLUE,
             bg_color: str = HEX_ICE, icon: str = "●") -> None:
    """Styled callout box for key findings."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Set paragraph shading
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shd)
    # Set borders via XML
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="{color}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)
    _make_run(p, f"  {icon}  ", bold=True, size=10, color=color)
    _make_run(p, text, bold=False, size=10, color=C.DARK_BLUE)


def _section_divider(doc: Document) -> None:
    """Thin decorative line between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{HEX_MID_BLUE}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════
# Headings & paragraphs
# ═══════════════════════════════════════════════════════════════════

def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        run.font.name = _CN_FONT or "Calibri"
        rPr = run._element.get_or_add_rPr()
        rPr.set(qn("w:eastAsia"), _CN_FONT or "")
        if level == 1:
            run.font.color.rgb = C.NAVY
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = C.DARK_BLUE
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = C.GRAY_DARK
            run.font.size = Pt(12)

def _para(doc: Document, text: str, bold: bool = False,
          size: int = 10, color: RGBColor = None,
          indent: float = 0, spacing_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _make_run(p, text, bold=bold, size=size, color=color or C.GRAY_DARK)

def _kv(doc: Document, key: str, value: str, indent: float = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _make_run(p, f"{key}：", bold=True, size=9.5, color=C.DARK_BLUE)
    _make_run(p, value, bold=False, size=9.5, color=C.GRAY_DARK)

def _bullet(doc: Document, text: str, color: RGBColor = None, indent: float = 0.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(indent)
    _make_run(p, text, size=9.5, color=color or C.GRAY_DARK)


# ═══════════════════════════════════════════════════════════════════
# Chart generation — professional styling
# ═══════════════════════════════════════════════════════════════════

_CHART_STYLE = {
    "axes.edgecolor": "#CCCCCC",
    "axes.grid": True,
    "grid.alpha": 0.3,
    "grid.color": "#E0E0E0",
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.titleweight": "bold",
    "axes.titlecolor": "#0B1F3F",
    "axes.titlelocation": "left",
}
plt.rcParams.update(_CHART_STYLE)

def _save_chart(fig: plt.Figure, tmp_dir: str, name: str) -> str:
    path = os.path.join(tmp_dir, f"{name}.png")
    fig.savefig(path, dpi=250, bbox_inches="tight", facecolor="white",
                edgecolor="none")
    plt.close(fig)
    return path

def _chart_pv(pv: list[float], title: str = "光伏典型日发电曲线") -> plt.Figure:
    """PV generation: filled area + peak label."""
    hours = list(range(24))
    pv24 = pv[:24]
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    ax.fill_between(hours, pv24, alpha=0.15, color=PLOT_ORANGE)
    ax.plot(hours, pv24, color=PLOT_ORANGE, linewidth=2.2)
    ax.set_xlabel("小时", fontsize=9, color="#666666")
    ax.set_ylabel("功率 (kW)", fontsize=9, color="#666666")
    ax.set_title(title, fontsize=12, pad=15)
    ax.set_xticks(range(0, 25, 2)); ax.set_xlim(0, 23)
    peak = max(pv24) if pv24 else 0
    if peak > 0:
        ph = pv24.index(peak)
        ax.annotate(f"{peak:,.0f} kW", xy=(ph, peak),
                    xytext=(ph + 2.5, peak * 0.92),
                    arrowprops=dict(arrowstyle="->", color=PLOT_ORANGE, lw=1.2),
                    fontsize=8, color=PLOT_ORANGE, fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.2", fc="white", ec=PLOT_ORANGE, alpha=0.8))
    fig.tight_layout()
    return fig

def _chart_load(load: list[float], title: str = "典型日负荷曲线") -> plt.Figure:
    """Load profile with peak annotation."""
    hours = list(range(24))
    ld24 = load[:24]
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    ax.fill_between(hours, ld24, alpha=0.1, color=PLOT_BLUE)
    ax.plot(hours, ld24, color=PLOT_BLUE, linewidth=2.2)
    ax.set_xlabel("小时", fontsize=9, color="#666666")
    ax.set_ylabel("负荷 (kW)", fontsize=9, color="#666666")
    ax.set_title(title, fontsize=12, pad=15)
    ax.set_xticks(range(0, 25, 2)); ax.set_xlim(0, 23)
    peak = max(ld24) if ld24 else 0
    if peak > 0:
        ph = ld24.index(peak)
        ax.annotate(f"{peak:,.0f} kW", xy=(ph, peak),
                    xytext=(ph + 2.5, peak * 0.92),
                    arrowprops=dict(arrowstyle="->", color=PLOT_BLUE, lw=1.2),
                    fontsize=8, color=PLOT_BLUE, fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.2", fc="white", ec=PLOT_BLUE, alpha=0.8))
    fig.tight_layout()
    return fig

def _chart_dispatch(hours: list[int], charge: list[float],
                    discharge: list[float], soc: list[float] | None = None,
                    title: str = "储能系统典型日运行曲线") -> plt.Figure:
    """Storage dispatch: charge/discharge bars + SOC line."""
    fig, ax1 = plt.subplots(figsize=(7.5, 3.5))
    w = 0.35
    ax1.bar([h - w/2 for h in hours], charge, w, color=PLOT_GREEN,
            alpha=0.8, label="充电功率")
    ax1.bar([h + w/2 for h in hours], [-d for d in discharge], w,
            color=PLOT_RED, alpha=0.8, label="放电功率")
    ax1.set_xlabel("小时", fontsize=9, color="#666666")
    ax1.set_ylabel("功率 (kW)", fontsize=9, color="#666666")
    ax1.set_xticks(range(0, 25, 2)); ax1.set_xlim(-0.5, 23.5)
    ax1.axhline(y=0, color="#999999", linewidth=0.5)
    ax1.legend(loc="upper left", fontsize=7, framealpha=0.9)

    if soc and any(abs(v) > 0.001 for v in soc):
        ax2 = ax1.twinx()
        ax2.plot(hours, [v * 100 for v in soc], color=PLOT_ORANGE,
                 linewidth=2, linestyle="--", marker="s", markersize=3,
                 label="SOC", zorder=5)
        ax2.set_ylabel("SOC (%)", fontsize=9, color=PLOT_ORANGE)
        ax2.set_ylim(0, 105)
        ax2.legend(loc="upper right", fontsize=7, framealpha=0.9)

    ax1.set_title(title, fontsize=12, pad=15)
    fig.tight_layout()
    return fig

def _chart_pie(labels: list[str], values: list[float],
               title: str) -> plt.Figure:
    """Donut-style pie chart."""
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    colors = PLOT_COLORS[:len(labels)]
    wedges, texts, autotexts = ax.pie(
        values, labels=None, autopct="%1.1f%%",
        startangle=90, colors=colors,
        pctdistance=0.82,
        wedgeprops=dict(width=0.35, edgecolor="white", linewidth=1.5))
    for at in autotexts: at.set_fontsize(7.5)
    ax.set_title(title, fontsize=12, pad=18)
    ax.legend(wedges, [f"  {l}" for l in labels],
              loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8,
              frameon=False)
    fig.tight_layout()
    return fig

def _chart_cashflow(years: list[int], cf: list[float],
                    npv_v: float | None = None) -> plt.Figure:
    """Cashflow: colored bars + cumulative line + zero line."""
    fig, ax = plt.subplots(figsize=(7.5, 3.8))
    colors = [PLOT_GREEN if v >= 0 else PLOT_RED for v in cf]
    ax.bar(years, cf, color=colors, width=0.55, edgecolor="white", linewidth=0.5)
    cumsum = np.cumsum(cf)
    ax.plot(years, cumsum, color=PLOT_BLUE, linewidth=2.2,
            marker="o", markersize=5, zorder=5, label="累计净现金流")
    ax.axhline(y=0, color="#AAAAAA", linewidth=0.8, linestyle="-")
    ax.set_xlabel("运营年份", fontsize=9, color="#666666")
    ax.set_ylabel("现金流 (元)", fontsize=9, color="#666666")
    ax.set_title("项目全生命周期现金流分析", fontsize=12, pad=15)
    ax.legend(fontsize=8, framealpha=0.9)
    if npv_v is not None:
        ax.text(0.97, 0.95, f"NPV = {npv_v:,.0f} 元",
                transform=ax.transAxes, fontsize=10, fontweight="bold",
                va="top", ha="right",
                bbox=dict(boxstyle="round,pad=0.4",
                          facecolor="#DFF0E3" if npv_v > 0 else "#F9E3E0",
                          edgecolor=PLOT_GREEN if npv_v > 0 else PLOT_RED,
                          alpha=0.9))
    fig.tight_layout()
    return fig

def _chart_sensitivity(factors: list[str], impacts: list[float]) -> plt.Figure:
    """Horizontal bar chart for sensitivity."""
    fig, ax = plt.subplots(figsize=(7, 3))
    colors_s = [PLOT_GREEN if v >= 0 else PLOT_RED for v in impacts]
    ax.barh(factors, impacts, color=colors_s, height=0.5, edgecolor="white")
    ax.set_xlabel("IRR 变化 (%)", fontsize=9, color="#666666")
    ax.set_title("IRR 敏感性分析", fontsize=12, pad=12)
    ax.axvline(x=0, color="#AAAAAA", linewidth=0.8)
    for i, (f, v) in enumerate(zip(factors, impacts)):
        ax.text(v + (0.001 if v >= 0 else -0.001), i,
                f"{v:+.2f}%", va="center",
                ha="left" if v >= 0 else "right", fontsize=8,
                color=PLOT_GREEN if v >= 0 else PLOT_RED)
    fig.tight_layout()
    return fig

def _chart_load_duration(load: list[float],
                        title: str = "负荷持续曲线") -> plt.Figure:
    """Load duration curve: sorted descending, with baseline/peak annotation."""
    sorted_load = sorted(load[:8760] if len(load) > 24 else load, reverse=True)
    hours = list(range(len(sorted_load)))
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    ax.fill_between(hours, sorted_load, alpha=0.12, color=PLOT_BLUE)
    ax.plot(hours, sorted_load, color=PLOT_BLUE, linewidth=1.8)
    p90 = sorted_load[int(len(sorted_load) * 0.1)] if sorted_load else 0
    p50 = sorted_load[int(len(sorted_load) * 0.5)] if sorted_load else 0
    ax.axhline(y=p90, color=PLOT_RED, linewidth=0.8, linestyle="--", alpha=0.6)
    ax.text(len(sorted_load) * 0.95, p90, f"P90: {p90:,.0f} kW",
            fontsize=7, color=PLOT_RED, va="bottom")
    ax.set_xlabel("小时", fontsize=9, color="#666666")
    ax.set_ylabel("负荷 (kW)", fontsize=9, color="#666666")
    ax.set_title(title, fontsize=12, pad=15)
    ax.set_xlim(0, len(sorted_load))
    fig.tight_layout()
    return fig

def _chart_pv_load_overlay(pv: list[float], load: list[float],
                           title: str = "光伏发电-负荷匹配曲线") -> plt.Figure:
    """Overlay PV and load on same chart to visualize self-consumption."""
    hours = list(range(24))
    pv24 = pv[:24]; ld24 = load[:24]
    surplus = [max(p - l, 0) for p, l in zip(pv24, ld24)]
    deficit = [max(l - p, 0) for l, p in zip(ld24, pv24)]
    fig, ax = plt.subplots(figsize=(7.5, 3.5))
    ax.fill_between(hours, ld24, alpha=0.2, color=PLOT_BLUE, label="负荷")
    ax.plot(hours, ld24, color=PLOT_BLUE, linewidth=2.2)
    ax.plot(hours, pv24, color=PLOT_ORANGE, linewidth=2.2, label="光伏")
    ax.fill_between(hours, surplus, alpha=0.25, color=PLOT_GREEN, label="余电上网")
    ax.fill_between(hours, deficit, alpha=0.15, color=PLOT_RED, label="电网取电")
    ax.set_xlabel("小时", fontsize=9, color="#666666")
    ax.set_ylabel("功率 (kW)", fontsize=9, color="#666666")
    ax.set_title(title, fontsize=12, pad=15)
    ax.set_xticks(range(0, 25, 2)); ax.set_xlim(0, 23)
    ax.legend(fontsize=7, ncol=2, framealpha=0.9, loc="upper left")
    fig.tight_layout()
    return fig

def _chart_revenue_pie(fin: dict[str, Any],
                       title: str = "年收益构成") -> plt.Figure | None:
    """Revenue composition donut chart."""
    items = [
        ("辅助服务", fin.get("annual_ancillary_service_revenue")),
        ("需求响应", fin.get("annual_demand_response_revenue")),
        ("外送收益", fin.get("annual_export_revenue")),
    ]
    labels = [l for l, v in items if v and v > 0]
    values = [v for _, v in items if v and v > 0]
    savings = fin.get("annual_savings_or_revenue")
    if savings and savings > 0:
        labels.insert(0, "综合节费"); values.insert(0, savings)
    if len(labels) < 2:
        return None
    fig, ax = plt.subplots(figsize=(5, 4.5))
    colors = PLOT_COLORS[:len(labels)]
    wedges, texts, autotexts = ax.pie(
        values, labels=None, autopct="%1.1f%%",
        startangle=90, colors=colors,
        pctdistance=0.82,
        wedgeprops=dict(width=0.35, edgecolor="white", linewidth=1.5))
    for at in autotexts: at.set_fontsize(8)
    ax.set_title(title, fontsize=12, pad=18)
    ax.legend(wedges, [f"  {l} ({v:,.0f})" for l, v in zip(labels, values)],
              loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8, frameon=False)
    fig.tight_layout()
    return fig

def _chart_generation_monthly(annual_series: list[float],
                               title: str = "月度发电量分布") -> plt.Figure | None:
    """Aggregate 8760 series into monthly generation bars."""
    if len(annual_series) < 8760:
        return None
    days_in_month = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
    monthly = []
    idx = 0
    for days in days_in_month:
        m = 0.0
        for d in range(days):
            for h in range(24):
                if idx < len(annual_series):
                    m += annual_series[idx]; idx += 1
        monthly.append(m / 1000)  # kWh → MWh
    months = ["1月","2月","3月","4月","5月","6月","7月","8月","9月","10月","11月","12月"]
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    bars = ax.bar(range(12), monthly, color=PLOT_ORANGE, width=0.6, edgecolor="white",
                  alpha=0.85)
    for bar, v in zip(bars, monthly):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height(),
                f"{v:,.0f}", ha="center", va="bottom", fontsize=7, color="#666666")
    ax.set_xticks(range(12)); ax.set_xticklabels(months, fontsize=8)
    ax.set_ylabel("发电量 (MWh)", fontsize=9, color="#666666")
    ax.set_title(title, fontsize=12, pad=15)
    fig.tight_layout()
    return fig

def _chart_monthly(months: list[str], charge_mwh: list[float],
                   discharge_mwh: list[float]) -> plt.Figure:
    """Grouped bar for monthly charge/discharge."""
    fig, ax = plt.subplots(figsize=(7.5, 3.5))
    x = np.arange(len(months))
    w = 0.32
    ax.bar(x - w, charge_mwh, w, color=PLOT_GREEN, alpha=0.8, label="充电")
    ax.bar(x, discharge_mwh, w, color=PLOT_RED, alpha=0.8, label="放电")
    ax.set_xticks(x); ax.set_xticklabels(months, fontsize=7.5)
    ax.set_ylabel("电量 (MWh)", fontsize=9, color="#666666")
    ax.set_title("月度储能运行电量", fontsize=12, pad=15)
    ax.legend(fontsize=8, framealpha=0.9)
    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════

def _build_cover(doc: Document, project_name: str, province: str,
                 detail_label: str) -> None:
    """Professional cover with color block design."""

    # ── Top accent bar ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="48" w:space="0" w:color="{HEX_NAVY}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)

    for _ in range(5):
        doc.add_paragraph()

    # ── Confidential mark ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _make_run(p, "  CONFIDENTIAL  ", bold=True, size=8, color=C.WHITE)
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_MID_BLUE}" w:val="clear"/>')
    pPr.append(shd)

    doc.add_paragraph()

    # ── Report type tag ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    _make_run(p, "新能源储能项目", size=12, color=C.MID_BLUE)

    # ── Main title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    _make_run(p, "可行性研究", size=30, bold=True, color=C.NAVY)
    _make_run(p, "报告", size=30, bold=True, color=C.NAVY)

    # ── Subtitle ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    _make_run(p, f"—— {project_name} ——", size=14, color=C.MID_BLUE)

    # ── Decorative line ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{HEX_MID_BLUE}"/>'
        f'</w:pBdr>')
    pPr.append(pBdr)

    # ── Info grid ──
    info = [
        ("委托单位", "________________________"),
        ("项目地点", province),
        ("方案类型", detail_label),
        ("报告编号", f"ESA-{datetime.now().strftime('%Y%m%d')}-{abs(hash(project_name)) % 10000:04d}"),
        ("编制日期", datetime.now().strftime("%Y年%m月%d日")),
        ("版    本", "V1.0 (初版)"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _make_run(p, f"  {label}    ", size=11, color=C.GRAY_DARK, bold=True)
        _make_run(p, value, size=11, color=C.GRAY_DARK)

    for _ in range(6):
        doc.add_paragraph()

    # ── Footer bar ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(p, "Energy Solution Agent  ·  AI 辅助分析系统", size=9, color=C.GRAY)

    # New section with proper page size
    new_sec = doc.add_section()
    new_sec.page_width = Cm(21)
    new_sec.page_height = Cm(29.7)
    new_sec.top_margin = Cm(2.2)
    new_sec.bottom_margin = Cm(2.2)
    new_sec.left_margin = Cm(2.5)
    new_sec.right_margin = Cm(2.5)


# ═══════════════════════════════════════════════════════════════════
# MAIN BUILDER
# ═══════════════════════════════════════════════════════════════════

def build_docx_report(output: dict[str, Any],
                      diagnostics: dict[str, Any],
                      filepath: str | Path) -> str:
    """Generate a professional design-institute DOCX report."""

    doc = Document()
    _setup_document(doc)
    tmp = tempfile.mkdtemp(prefix="esa_rpt_")

    # ── Data extraction ──
    S  = output.get("project_summary", {})
    sol = output.get("recommended_solution", {})
    sim = output.get("simulation_results", {})
    disp = output.get("dispatch_results", {})
    res = output.get("resource_results", {})
    fin = output.get("financial_results", {})
    carb = output.get("carbon_results", {})
    mkt = output.get("market_and_settlement", {})
    sens = output.get("sensitivity_results", [])
    des = output.get("design_and_interconnection", {})
    alt = output.get("alternative_solutions", [])
    qual = output.get("data_quality_results", {})
    gaps = output.get("data_gaps", [])
    risks = output.get("risks", [])

    scenario = (S.get("scenario_type") or "").lower()
    detail_label = S.get("scenario_detail_label") or S.get("scenario_type") or "未知"
    proj_name = S.get("project_name") or "未命名项目"
    province = S.get("province") or "—"
    op_mode = S.get("operation_mode") or "—"

    has_pv = bool(sol.get("pv_mwp") or res.get("pv_p50_generation_mwh"))
    has_wind = bool(sol.get("wind_mw"))
    has_charging = bool(sol.get("charging_capacity_mw") or sim.get("annual_charging_energy_mwh"))
    has_thermal = bool(sol.get("cooling_capacity_rt") or sol.get("heating_capacity_mwth"))
    has_market = mkt.get("market_mode") not in (None, "", "none")

    # ── Build pages ──
    _build_cover(doc, proj_name, province, detail_label)
    doc.add_page_break()

    # TOC
    _heading(doc, "目  录", 1)
    _para(doc, "（请在 Word 中右键此处 → 更新域，自动生成目录）", size=9, color=C.GRAY)
    _add_toc(doc)
    doc.add_page_break()

    # 1. Executive Summary
    _heading(doc, "1  执行摘要", 1)
    _build_exec_summary(doc, output, S, sol, fin, sim, carb, disp, res, tmp)

    # 2. Project Overview
    _heading(doc, "2  项目概况", 1)
    _build_project_info(doc, S, sol, des, qual, diagnostics, gaps, risks)

    # 3. Resource Assessment
    if has_pv or has_wind:
        _heading(doc, "3  资源评估", 1)
        _build_resource(doc, res, sim, has_pv, has_wind, tmp)

    # 4. Load Analysis
    _heading(doc, "4  负荷分析", 1)
    _build_load(doc, sim, disp, tmp)

    # 5. Storage System
    _heading(doc, "5  储能系统", 1)
    _build_storage(doc, sol, disp, tmp)

    # Scenario sections
    _sec = 6
    if has_charging:
        _heading(doc, f"{_sec}  充电设施", 1); _sec += 1
        _build_charging(doc, disp, sim)
    if has_thermal:
        _heading(doc, f"{_sec}  暖通空调", 1); _sec += 1
        _build_thermal(doc, disp, sim)
    if has_market:
        _heading(doc, f"{_sec}  电力市场与交易", 1); _sec += 1
        _build_market(doc, mkt)

    # Financial
    _heading(doc, f"{_sec}  财务分析", 1); _sec += 1
    _build_financial(doc, fin, sol, sim, alt, tmp)

    # Carbon
    _heading(doc, f"{_sec}  碳减排分析", 1); _sec += 1
    _build_carbon(doc, carb)

    # Sensitivity
    if sens:
        _heading(doc, f"{_sec}  敏感性分析", 1)
        _build_sensitivity(doc, sens, tmp)

    # Conclusion
    _heading(doc, f"{_sec + 1}  结论与建议", 1)
    _build_conclusion(doc, output, fin, carb)

    # Headers & footers
    _add_headers_footers(doc, proj_name)

    out_path = Path(filepath)
    doc.save(str(out_path))
    return str(out_path.resolve())


# ═══════════════════════════════════════════════════════════════════
# SECTION BUILDERS
# ═══════════════════════════════════════════════════════════════════

def _build_exec_summary(doc, output, S, sol, fin, sim, carb, disp, res, tmp):
    """Executive summary with KPI dashboard, verdict, and energy chart."""
    irr = fin.get("irr")
    npv = fin.get("npv")
    payback = fin.get("payback_years")
    abate = carb.get("annual_reduction_tco2e")
    stg_mwh = sol.get("storage_energy_mwh") or 0
    stg_mw = sol.get("storage_power_mw") or 0
    pv_mw = sol.get("pv_mwp") or 0
    pv_gen = sim.get("annual_pv_generation_mwh") or 0
    wind_mw = sol.get("wind_mw") or 0
    wind_gen = sim.get("annual_wind_generation_mwh") or 0

    _para(doc, "本报告对项目技术可行性、经济合理性及碳减排效益进行了全面分析。", size=10)

    # ── KPI Dashboard ──
    _heading(doc, "1.1  核心指标", 2)

    # Format KPI values
    irr_str = f"{irr * 100:.2f}%" if irr is not None else "N/A"
    npv_str = f"{npv / 1e4:,.1f} 万" if npv is not None else "N/A"
    payback_str = f"{payback:.1f} 年" if payback is not None else "N/A"
    abate_str = f"{abate:,.0f} t" if abate else "—"

    irr_color = C.GREEN if (irr and irr >= 0.08) else (C.RED if irr and irr < 0 else C.ORANGE)
    npv_color = C.GREEN if (npv and npv > 0) else (C.RED if npv and npv <= 0 else C.ORANGE)

    cards = [
        ("IRR 内部收益率", irr_str, "基准 ≥ 8%", irr_color),
        ("NPV 净现值", npv_str, "基准 > 0", npv_color),
        ("静态回收期", payback_str, "含建设期", C.MID_BLUE),
        ("年减排量", abate_str, "CO₂ 当量", C.GREEN),
    ]
    _add_kpi_cards(doc, cards)

    # ── Verdict callout ──
    conclusion = S.get("conclusion_level", "pre_feasibility")
    cmap = {"feasibility": "可行性研究", "pre_feasibility": "预可行性研究", "concept": "概念设计"}
    _callout(doc, f"分析等级：{cmap.get(conclusion, conclusion)}。"
             f"配置容量 {pv_mw} MWp / {wind_mw} MW / {stg_mw} MW，储能规模 {stg_mwh} MWh。",
             color=C.MID_BLUE, bg_color=HEX_ICE)

    # ── Key metrics table ──
    _heading(doc, "1.2  关键参数", 2)
    rows = [
        ["装机容量", f"{pv_mw} MWp (光伏) / {wind_mw} MW (风电)"],
        ["储能规模", f"{stg_mwh} MWh / {stg_mw} MW"],
        ["年光伏发电量", f"{_fmt(pv_gen)} MWh"],
        ["年风电发电量", f"{_fmt(wind_gen)} MWh"],
        ["年储能放电量", f"{_fmt(sim.get('annual_storage_discharge_mwh'))} MWh"],
        ["年等效满循环", f"{_fmt(disp.get('storage_equivalent_full_cycles_per_year'))} 次"],
        ["估算储能寿命", f"{_fmt(disp.get('storage_life_years_estimate'))} 年"],
        ["年购电量", f"{_fmt(sim.get('annual_grid_purchase_mwh'))} MWh"],
        ["年外送电量", f"{_fmt(sim.get('annual_export_mwh'))} MWh"],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5], first_col_bold=True)

    # ── Energy mix pie + PV-Load overlay ──
    labels_p, values_p = [], []
    if pv_gen > 0: labels_p.append("光伏"); values_p.append(pv_gen)
    if wind_gen > 0: labels_p.append("风电"); values_p.append(wind_gen)
    grid_p = sim.get("annual_grid_purchase_mwh") or 0
    if grid_p > 0: labels_p.append("电网购电"); values_p.append(grid_p)
    stg_d = sim.get("annual_storage_discharge_mwh") or 0
    if stg_d > 0: labels_p.append("储能放电"); values_p.append(stg_d)
    if len(labels_p) >= 2:
        fig = _chart_pie(labels_p, values_p, "年度能源供给结构")
        path = _save_chart(fig, tmp, "energy_mix")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(13))
        _para(doc, "图 1-1  年度能源供给结构", size=8, color=C.GRAY)

    # PV-Load matching chart
    pv24_summary = (sim.get("pv_hourly_profile_kw") or res.get("pv_hourly_profile_kw") or [])[:24]
    load24_summary = (disp.get("load_profile_kw") or sim.get("load_hourly_profile_kw") or [])[:24]
    if pv24_summary and max(pv24_summary) > 0 and load24_summary and max(load24_summary) > 0:
        fig = _chart_pv_load_overlay(pv24_summary, load24_summary)
        path = _save_chart(fig, tmp, "pv_load_match")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 1-2  光伏发电与负荷匹配曲线", size=8, color=C.GRAY)

    doc.add_page_break()


def _build_project_info(doc, S, sol, des, qual, diagnostics, gaps, risks):
    """Project overview."""
    rows = [
        ["项目名称", S.get("project_name") or "—"],
        ["场景类型", S.get("scenario_detail_label") or S.get("scenario_type") or "—"],
        ["省份", S.get("province") or "—"],
        ["运行模式", S.get("operation_mode") or "—"],
        ["分析模式", S.get("analysis_mode") or "—"],
        ["数据完整度", diagnostics.get("data_completeness_grade") or "—"],
        ["数据质量", f"{qual.get('score', '—')} ({qual.get('level', '—')})"],
    ]
    _add_table(doc, ["项目", "内容"], rows, col_widths=[4.5, 12.5], first_col_bold=True)

    if des.get("recommended_voltage_level_kv"):
        _heading(doc, "2.1  接入方案", 2)
        _kv(doc, "推荐电压等级", f"{des.get('recommended_voltage_level_kv')} kV")
        _kv(doc, "接入方式", des.get("recommended_connection_mode") or "—")
        for n in des.get("primary_system_notes") or []:
            _bullet(doc, f"电气一次：{n}")
        for n in des.get("secondary_system_notes") or []:
            _bullet(doc, f"电气二次：{n}")
        for s in des.get("required_studies") or []:
            _bullet(doc, f"专题：{s}")
        for a in des.get("required_approvals") or []:
            _bullet(doc, f"批复：{a}")

    if gaps or risks:
        _heading(doc, "2.2  风险与缺口", 2)
        for g in gaps:
            _bullet(doc, f"数据缺口：{g}", color=C.ORANGE)
        for r in risks:
            _bullet(doc, f"风险：{r}", color=C.RED)


def _build_resource(doc, res, sim, has_pv, has_wind, tmp):
    """PV & Wind resource assessment."""
    if has_pv:
        _heading(doc, "3.1  光伏资源", 2)
        pv_rows = [
            ["资源口径", f"{res.get('pv_resource_basis') or '—'} / {res.get('pv_resource_accuracy') or '—'}"],
            ["P50 年发电量", f"{_fmt(res.get('pv_p50_generation_mwh'))} MWh"],
            ["P90 年发电量", f"{_fmt(res.get('pv_p90_generation_mwh'))} MWh"],
            ["综合 PR", _fmt(res.get("pv_pr_effective"))],
            ["倾角", f"{_fmt(res.get('pv_effective_tilt_deg'), 0)}° (推荐 {_fmt(res.get('pv_recommended_tilt_deg'), 0)}°)"],
            ["倾角因子", _fmt(res.get("pv_tilt_factor"))],
            ["方位角因子", _fmt(res.get("pv_azimuth_factor"))],
            ["跟踪因子", _fmt(res.get("pv_tracking_factor"))],
            ["温度因子", _fmt(res.get("pv_temperature_factor"))],
            ["双面增益", _fmt(res.get("pv_bifacial_gain"))],
            ["遮挡因子", _fmt(res.get("pv_shading_factor"))],
        ]
        _add_table(doc, ["参数", "数值"], pv_rows, col_widths=[5, 12])
        pv24 = (sim.get("pv_hourly_profile_kw") or res.get("pv_hourly_profile_kw") or [])[:24]
        if pv24 and max(pv24) > 0:
            fig = _chart_pv(pv24)
            path = _save_chart(fig, tmp, "pv_profile")
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Cm(15))
            _para(doc, "图 3-1  光伏典型日发电曲线", size=8, color=C.GRAY)

        # Monthly PV generation
        pv_annual = sim.get("pv_annual_series_kw") or []
        if pv_annual and len(pv_annual) >= 8760 and max(pv_annual) > 0:
            fig = _chart_generation_monthly(pv_annual, "光伏月度发电量分布")
            if fig:
                path = _save_chart(fig, tmp, "pv_monthly")
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Cm(15))
                _para(doc, "图 3-2  光伏月度发电量分布", size=8, color=C.GRAY)

    if has_wind:
        _heading(doc, "3.2  风能资源", 2)
        wind_rows = [
            ["资源口径", f"{res.get('wind_resource_basis') or '—'} / {res.get('wind_resource_accuracy') or '—'}"],
            ["P50", f"{_fmt(res.get('wind_p50_generation_mwh'))} MWh"],
            ["P90", f"{_fmt(res.get('wind_p90_generation_mwh'))} MWh"],
            ["功率曲线", _fmt(res.get("wind_power_curve_used"))],
            ["净折减因子", _fmt(res.get("wind_net_factor"))],
        ]
        _add_table(doc, ["参数", "数值"], wind_rows, col_widths=[5, 12])


def _build_load(doc, sim, disp, tmp):
    """Load analysis."""
    rows = [
        ["年网购电量", f"{_fmt(sim.get('annual_grid_purchase_mwh'))} MWh"],
        ["年外送电量", f"{_fmt(sim.get('annual_export_mwh'))} MWh"],
        ["基线峰值购电", f"{_fmt(disp.get('baseline_peak_grid_kw'))} kW"],
        ["储能后峰值购电", f"{_fmt(disp.get('post_storage_peak_grid_kw'))} kW"],
        ["削峰量", f"{_fmt(disp.get('estimated_peak_reduction_kw'))} kW"],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5])

    load24 = (disp.get("load_profile_kw") or sim.get("load_hourly_profile_kw") or [])[:24]
    if load24 and max(load24) > 0:
        fig = _chart_load(load24)
        path = _save_chart(fig, tmp, "load_profile")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 4-1  典型日负荷曲线", size=8, color=C.GRAY)

    # Load duration curve
    load_series = disp.get("load_profile_kw") or sim.get("load_hourly_profile_kw") or []
    if load_series and max(load_series) > 0:
        fig = _chart_load_duration(load_series, "负荷持续曲线")
        path = _save_chart(fig, tmp, "load_duration")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 4-2  负荷持续曲线（P90 红线）", size=8, color=C.GRAY)
    doc.add_page_break()


def _build_storage(doc, sol, disp, tmp):
    """Storage system design."""
    stg_mwh = sol.get("storage_energy_mwh") or 0
    stg_mw = sol.get("storage_power_mw") or 0
    raw_mwh = sol.get("raw_storage_energy_mwh") or 0
    raw_mw = sol.get("raw_storage_power_mw") or 0

    _heading(doc, "5.1  容量配置", 2)
    rows = [
        ["推荐容量", f"{stg_mw} MW / {stg_mwh} MWh"],
        ["计算容量", f"{raw_mw} MW / {raw_mwh} MWh" if raw_mwh else "—"],
        ["充放电时长", f"{stg_mwh / max(stg_mw, 0.001):.1f} h"],
        ["策略模式", disp.get("storage_strategy_mode") or "—"],
        ["测算口径", disp.get("storage_sizing_basis") or "—"],
    ]
    _add_table(doc, ["参数", "数值"], rows, col_widths=[5.5, 11.5])

    _heading(doc, "5.2  运行指标", 2)
    op_rows = [
        ["日循环次数", _fmt(disp.get("daily_storage_cycles"))],
        ["年等效满循环", _fmt(disp.get("storage_equivalent_full_cycles_per_year"))],
        ["年吞吐量", f"{_fmt(disp.get('storage_annual_throughput_mwh'))} MWh"],
        ["估算寿命", f"{_fmt(disp.get('storage_life_years_estimate'))} 年"],
        ["往返效率", _pct(disp.get("storage_effective_round_trip_efficiency"))],
        ["年衰减率", _pct(disp.get("storage_degradation_per_year"))],
        ["寿命末容量比", _pct(disp.get("storage_end_of_life_capacity_ratio"))],
        ["绿电充电比", _pct(disp.get("storage_charge_from_renewables_ratio"))],
        ["预留 SOC", _pct(disp.get("storage_reserved_soc_ratio"))],
        ["保供 SOC", _pct(disp.get("storage_backup_soc_ratio"))],
    ]
    _add_table(doc, ["指标", "数值"], op_rows, col_widths=[5.5, 11.5])

    # Dispatch chart
    dp = disp.get("dispatch_series_kw") or []
    if dp and len(dp) >= 24:
        hrs = list(range(24))
        ch, dch, soc_hr = [], [], []
        soc_raw = disp.get("soc_series") or []
        soc_m = max(max(soc_raw[:24], key=lambda x: max(x, 0)), 0.01) if soc_raw else 1
        for h in range(24):
            v = dp[h] if h < len(dp) else 0
            ch.append(max(v, 0)); dch.append(abs(min(v, 0)))
            soc_hr.append((soc_raw[h] / soc_m) if soc_raw and h < len(soc_raw) else 0.5)
        fig = _chart_dispatch(hrs, ch, dch, soc_hr)
        path = _save_chart(fig, tmp, "dispatch")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 5-1  储能典型日充放电曲线", size=8, color=C.GRAY)

    # Monthly breakdown
    monthly = disp.get("monthly_storage_revenue_breakdown") or []
    if monthly:
        _heading(doc, "5.3  月度运行", 2)
        m_rows = [[f"{m.get('month','')}月", _fmt(m.get("charge_mwh")),
                   _fmt(m.get("discharge_mwh")), _fmt(m.get("gross_margin"))]
                  for m in monthly]
        _add_table(doc, ["月份", "充电(MWh)", "放电(MWh)", "毛收益(元)"],
                   m_rows, col_widths=[3, 4.5, 4.5, 5])

        # Monthly chart
        months = [f"{m.get('month','')}" for m in monthly]
        ch_m = [float(m.get("charge_mwh") or 0) for m in monthly]
        dch_m = [float(m.get("discharge_mwh") or 0) for m in monthly]
        if any(v > 0 for v in ch_m + dch_m):
            fig = _chart_monthly(months, ch_m, dch_m)
            path = _save_chart(fig, tmp, "monthly_dispatch")
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Cm(15))
            _para(doc, "图 5-2  月度储能运行电量", size=8, color=C.GRAY)
    doc.add_page_break()


def _build_charging(doc, disp, sim):
    """Charging station section."""
    rows = [
        ["年充电量", f"{_fmt(sim.get('annual_charging_energy_mwh'))} MWh"],
        ["充电峰值", f"{_fmt(disp.get('charging_peak_kw'))} kW"],
        ["利用率", _pct(disp.get("charging_utilization_ratio"))],
        ["排队风险", disp.get("charging_queue_risk") or "—"],
        ["排队指数", _fmt(disp.get("charging_queue_index"))],
        ["多样性系数", _fmt(disp.get("charging_diversity_factor"))],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5])
    seg = disp.get("charging_segment_summary") or []
    if seg:
        _heading(doc, "  充电结构", 2)
        seg_rows = [[s.get("vehicle_type", "—"),
                     f"{_fmt(s.get('daily_energy_kwh'))} kWh",
                     f"{_fmt(s.get('peak_kw'))} kW"] for s in seg]
        _add_table(doc, ["车型", "日充电量", "峰值功率"], seg_rows, col_widths=[5, 6, 6])


def _build_thermal(doc, disp, sim):
    """Thermal (HVAC) section."""
    rows = [
        ["年供冷量", f"{_fmt(sim.get('annual_cooling_energy_mwh'))} MWh"],
        ["年供热量", f"{_fmt(sim.get('annual_heating_energy_mwh'))} MWh"],
        ["冷峰值", f"{_fmt(disp.get('thermal_cooling_peak_kwth'))} kWth"],
        ["热峰值", f"{_fmt(disp.get('thermal_heating_peak_kwth'))} kWth"],
        ["制冷电耗峰值", f"{_fmt(disp.get('thermal_electric_peak_kw'))} kW"],
        ["锅炉燃料当量", f"{_fmt(disp.get('thermal_annual_boiler_fuel_equivalent_mwh'))} MWh"],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5])


def _build_market(doc, mkt):
    """Market & trading section."""
    rows = [
        ["市场模式", mkt.get("market_mode") or "—"],
        ["政策模式", mkt.get("market_policy_mode") or "—"],
        ["充/放电基准价", f"{_fmt(mkt.get('trading_charge_benchmark_price_per_kwh'))} / "
                        f"{_fmt(mkt.get('trading_discharge_benchmark_price_per_kwh'))} 元/kWh"],
        ["交易价差", f"{_fmt(mkt.get('trading_price_spread_per_kwh'))} 元/kWh"],
        ["波动指数", _fmt(mkt.get("trading_volatility_index"))],
        ["最优/最弱月份", f"{mkt.get('trading_best_month','—')} / {mkt.get('trading_worst_month','—')}"],
        ["规则生效字段", ", ".join(sorted(mkt.get("live_rule_effective_patch", {}).keys()))
                         if mkt.get("live_rule_effective_patch") else "—"],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5])

    if mkt.get("trading_execution_summary"):
        _callout(doc, mkt["trading_execution_summary"], color=C.MID_BLUE, bg_color=HEX_ICE)
    if mkt.get("cooptimization_execution_summary"):
        _heading(doc, "  协同优化", 2)
        _para(doc, mkt["cooptimization_execution_summary"], size=9.5)
    if mkt.get("spot_trading_cycle_summary"):
        _heading(doc, "  日内套利", 2)
        _para(doc, mkt["spot_trading_cycle_summary"], size=9.5)


def _build_financial(doc, fin, sol, sim, alt, tmp):
    """Financial analysis: CAPEX, OPEX, revenue, tax, KPIs, cashflow chart."""
    irr = fin.get("irr")
    npv = fin.get("npv")
    payback = fin.get("payback_years")

    # ── KPI cards ──
    irr_c = C.GREEN if (irr and irr >= 0.08) else (C.RED if irr and irr < 0.06 else C.ORANGE)
    npv_c = C.GREEN if (npv and npv > 0) else C.RED
    kpi_cards = [
        ("IRR", f"{irr * 100:.2f}%" if irr is not None else "N/A", "基准≥8%", irr_c),
        ("NPV", f"{npv / 1e4:,.1f}万" if npv is not None else "N/A", "基准>0", npv_c),
        ("回收期", f"{payback:.1f}年" if payback is not None else "N/A", "含建设期", C.MID_BLUE),
        ("LCOE", _fmt(fin.get("lcoe")) if fin.get("lcoe") else "N/A", "元/kWh", C.MID_BLUE),
    ]
    _add_kpi_cards(doc, kpi_cards)

    _heading(doc, "  投资与成本", 2)
    capex_rows = [
        ["系统总投资", f"{_fmt(fin.get('capex_total'))} 元"],
        ["年运维费", f"{_fmt(fin.get('opex_annual'))} 元/年"],
        ["运维递增率", _pct(fin.get("opex_escalation_rate"))],
        ["储能更换年份", _fmt(fin.get("storage_replacement_year"), 0) if fin.get("storage_replacement_year") else "—"],
        ["更换成本", f"{_fmt(fin.get('storage_replacement_cost'))} 元" if fin.get("storage_replacement_cost") else "—"],
    ]
    _add_table(doc, ["项目", "金额"], capex_rows, col_widths=[5.5, 11.5])

    _heading(doc, "  收益结构", 2)
    rev_rows = [
        ["年收益/节费", f"{_fmt(fin.get('annual_savings_or_revenue'))} 元"],
        ["电量电费", f"{_fmt(fin.get('annual_energy_charge_cost'))} 元"],
        ["需量电费", f"{_fmt(fin.get('annual_demand_charge_cost'))} 元"],
        ["辅助服务", f"{_fmt(fin.get('annual_ancillary_service_revenue'))} 元"],
        ["需求响应", f"{_fmt(fin.get('annual_demand_response_revenue'))} 元"],
        ["外送收益", f"{_fmt(fin.get('annual_export_revenue'))} 元"],
    ]
    _add_table(doc, ["项目", "金额"], rev_rows, col_widths=[5.5, 11.5])

    _heading(doc, "  税务", 2)
    tax_rows = [
        ["税模型", fin.get("tax_model") or "—"],
        ["年税费合计", f"{_fmt(fin.get('annual_tax_total'))} 元"],
        ["企业所得税", f"{_fmt(fin.get('annual_income_tax'))} 元"],
        ["增值税(应纳)", f"{_fmt(fin.get('annual_vat_payable'))} 元"],
        ["附加税", f"{_fmt(fin.get('annual_vat_surcharges_only'))} 元"],
        ["期初留抵", f"{_fmt(fin.get('initial_input_vat_credit'))} 元"],
    ]
    _add_table(doc, ["项目", "金额"], tax_rows, col_widths=[5.5, 11.5])

    # Revenue composition chart
    rev_fig = _chart_revenue_pie(fin)
    if rev_fig:
        path = _save_chart(rev_fig, tmp, "revenue_pie")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(12))
        _para(doc, "图 6-1  年收益构成", size=8, color=C.GRAY)

    _heading(doc, "  综合指标", 2)
    kpi_rows = [
        ["IRR", f"{_pct(irr)}" if irr is not None else "N/A",
         "✓ 达标" if irr and irr >= 0.08 else ("—" if irr is None else "✗ 未达标")],
        ["NPV", f"{_fmt(npv)} 元" if npv is not None else "N/A",
         "✓ 达标" if npv and npv > 0 else ("—" if npv is None else "✗ 未达标")],
        ["回收期", f"{_fmt(payback)} 年" if payback is not None else "N/A",
         "≤8 年为良好"],
        ["LCOE", f"{_fmt(fin.get('lcoe'))} 元/kWh" if fin.get("lcoe") else "—", "—"],
        ["减碳成本", f"{_fmt(fin.get('abatement_cost_per_tco2e'))} 元/tCO₂" if fin.get("abatement_cost_per_tco2e") else "—", "—"],
    ]
    _add_table(doc, ["指标", "数值", "评价"], kpi_rows, col_widths=[5, 6, 6])

    # Cashflow chart
    cf = fin.get("cashflow_series") or []
    if cf and len(cf) > 1:
        years = list(range(len(cf)))
        fig = _chart_cashflow(years, cf, npv)
        path = _save_chart(fig, tmp, "cashflow")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 6-2  全生命周期现金流分析", size=8, color=C.GRAY)

    # Alternative solutions
    if alt:
        _heading(doc, "  方案对比", 2)
        alt_rows = [[a.get("label", "—"),
                     f"{_fmt(a.get('storage_power_mw'))}MW/{_fmt(a.get('storage_energy_mwh'))}MWh",
                     _pct(a.get("irr")), _fmt(a.get("payback_years")), _fmt(a.get("npv"))]
                    for a in alt]
        _add_table(doc, ["方案", "容量", "IRR", "回收期", "NPV"],
                   alt_rows, col_widths=[3.5, 4, 3, 3, 3.5])
    doc.add_page_break()


def _build_carbon(doc, carb):
    """Carbon reduction analysis."""
    abate = carb.get("annual_reduction_tco2e")
    if abate and float(abate) > 0:
        _callout(doc, f"项目年减排 CO₂ 约 {_fmt(abate)} 吨，等效植树 {_fmt(float(abate) * 50)} 棵/年。",
                 color=C.GREEN, bg_color=HEX_LIGHT_GREEN, icon="✓")

    rows = [
        ["基线排放", f"{_fmt(carb.get('baseline_emissions_tco2e'))} tCO₂e"],
        ["项目后排放", f"{_fmt(carb.get('post_project_emissions_tco2e'))} tCO₂e"],
        ["年减排量", f"{_fmt(abate)} tCO₂e"],
        ["范围1", f"{_fmt(carb.get('scope1_reduction_tco2e'))} tCO₂e"],
        ["范围2", f"{_fmt(carb.get('scope2_reduction_tco2e'))} tCO₂e"],
        ["范围3", f"{_fmt(carb.get('scope3_reduction_tco2e'))} tCO₂e"],
        ["绿电覆盖率", _pct(carb.get("green_power_coverage_ratio"))],
        ["声明边界", carb.get("claim_boundary_summary") or "—"],
        ["减碳成本", f"{_fmt(carb.get('abatement_cost_per_tco2e'))} 元/tCO₂"],
    ]
    _add_table(doc, ["指标", "数值"], rows, col_widths=[5.5, 11.5])

    pbd = carb.get("carbon_path_breakdown") or []
    if pbd:
        _heading(doc, "  减排路径", 2)
        pr = [[p.get("path", "—"), f"{_fmt(p.get('reduction_tco2e'))} tCO₂e", _pct(p.get("share"))]
              for p in pbd]
        _add_table(doc, ["减排路径", "减排量", "占比"], pr, col_widths=[7, 5, 4])

    tmpl = carb.get("industry_template") or {}
    if tmpl:
        _heading(doc, "  行业建议", 2)
        for p in tmpl.get("major_paths", []):
            _bullet(doc, f"推荐：{p}")


def _build_sensitivity(doc, sens, tmp):
    """Sensitivity analysis."""
    sr = [[s.get("factor", "—"), _fmt(s.get("impact_on_annual_revenue")),
           _fmt(s.get("impact_on_irr"))] for s in sens]
    _add_table(doc, ["敏感因素", "年收益影响", "IRR 敏感度"], sr, col_widths=[6, 5.5, 5.5])

    factors = [s.get("factor", "") for s in sens]
    impacts = []
    for s in sens:
        v = s.get("impact_on_irr")
        try: impacts.append(float(v or 0))
        except: impacts.append(0)
    if factors and any(abs(v) > 0.001 for v in impacts):
        fig = _chart_sensitivity(factors, impacts)
        path = _save_chart(fig, tmp, "sensitivity")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(15))
        _para(doc, "图 6-2  IRR 敏感性分析", size=8, color=C.GRAY)


def _build_conclusion(doc, output, fin, carb):
    """Conclusion & recommendations."""
    irr, npv, abate = fin.get("irr"), fin.get("npv"), carb.get("annual_reduction_tco2e")

    if irr is not None and npv is not None:
        if irr >= 0.08 and npv > 0:
            verdict = "项目具备良好的经济可行性，推荐推进实施。"
            vc, bg = C.GREEN, HEX_LIGHT_GREEN
        elif irr >= 0.06:
            verdict = "项目经济性一般，建议优化配置或争取补贴后推进。"
            vc, bg = C.ORANGE, HEX_LIGHT_ORANGE
        else:
            verdict = "项目经济性偏低，建议重新评估或优化方案。"
            vc, bg = C.RED, HEX_LIGHT_RED
    else:
        verdict = "数据不足，建议补充完整参数后重新评估。"
        vc, bg = C.MID_BLUE, HEX_ICE

    if abate:
        verdict += f" 环境效益：年减排 CO₂ {_fmt(abate)} 吨。"
    _callout(doc, verdict, color=vc, bg_color=bg, icon="■")

    _heading(doc, "  下一步工作建议", 2)
    recs = [
        "补充8760点完整年负荷曲线数据，提高储能容量测算精度",
        "获取当地电网公司并网批复意见及接入系统设计方案",
        "确认储能设备品牌选型、技术参数及供应商报价（≥3家比选）",
        "核实省级最新峰谷电价政策及需量计费方式",
        "开展现场踏勘，确定设备布置方案及土建条件",
        "委托有资质的第三方编制施工图设计及消防专篇",
    ]
    for i, r in enumerate(recs, 1):
        _bullet(doc, f"{i}. {r}")

    _heading(doc, "  免责声明", 2)
    _para(doc, "本报告由 AI 辅助分析系统自动生成，仅供项目前期决策参考。最终投资决策应结合"
          "现场条件、设备实际报价、电网接入批复及专业设计院意见综合判断。储能系统实际运行"
          "效果受电价政策变动、电池衰减特性、负荷变化等多因素影响。", size=9, color=C.GRAY)


# ═══════════════════════════════════════════════════════════════════
# Document infrastructure
# ═══════════════════════════════════════════════════════════════════

def _setup_document(doc: Document) -> None:
    """Configure styles and defaults."""
    style = doc.styles["Normal"]
    style.font.name = _CN_FONT or "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.45
    style.paragraph_format.space_after = Pt(4)
    rPr = style.element.get_or_add_rPr()
    rPr.set(qn("w:eastAsia"), _CN_FONT or "")

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

def _add_headers_footers(doc: Document, project_name: str) -> None:
    """Header with project name + footer with page numbers and decorative line."""
    for sec in doc.sections:
        # ── Header ──
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        _make_run(hp, project_name, size=7.5, color=C.GRAY)
        _make_run(hp, "  |  ", size=7.5, color=C.GRAY_LIGHT)
        _make_run(hp, "工商业储能可行性研究报告", size=7.5, color=C.GRAY)

        # Bottom line
        hpPr = hp._element.get_or_add_pPr()
        hpBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{HEX_MID_BLUE}"/>'
            f'</w:pBdr>')
        hpPr.append(hpBdr)

        # ── Footer with page number ──
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)

        # Top decorative line
        fpPr = fp._element.get_or_add_pPr()
        fpBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{HEX_MID_BLUE}"/>'
            f'</w:pBdr>')
        fpPr.append(fpBdr)

        # Page X of Y field
        def _add_field(paragraph, code):
            r = paragraph.add_run()
            r._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
            r2 = paragraph.add_run()
            r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> {code} </w:instrText>'))
            r3 = paragraph.add_run()
            r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

        fp_run = fp.add_run("— ")
        fp_run.font.size = Pt(8)
        fp_run.font.color.rgb = C.GRAY
        _add_field(fp, "PAGE")
        fp_run2 = fp.add_run(" / ")
        fp_run2.font.size = Pt(8)
        fp_run2.font.color.rgb = C.GRAY
        _add_field(fp, "NUMPAGES")
        fp_run3 = fp.add_run(" —")
        fp_run3.font.size = Pt(8)
        fp_run3.font.color.rgb = C.GRAY


def _add_toc(doc: Document) -> None:
    """Insert Word TOC field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = p.add_run(); r._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = p.add_run(); r2._element.append(
        parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
    r3 = p.add_run(); r3._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    r4 = p.add_run("（在 Word 中右键 → 更新域 生成目录）")
    r4.font.size = Pt(10); r4.font.color.rgb = C.GRAY
    r5 = p.add_run(); r5._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
