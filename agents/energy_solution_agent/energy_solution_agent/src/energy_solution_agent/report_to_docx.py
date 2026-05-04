"""专家级 Word 报告生成器"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def _set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def _add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _set_cell_shading(cell, '2F5496')
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255,255,255)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 0:
                _set_cell_shading(cell, 'D6E4F0')
    return t

from .analysis_narrative import generate_narrative
from .report_charts import build_report_charts

def build_docx(result_json: dict) -> str:
    out = result_json
    f = out.get('financial_results', {})
    s = out.get('recommended_solution', {})
    d = out.get('dispatch_results', {})
    sim = out.get('simulation_results', {})
    carbon = out.get('carbon_results', {})
    sens = out.get('sensitivity_results', [])
    market = out.get('market_and_settlement', {})
    proj = out.get('project_summary', {})
    risks = out.get('risks', [])

    doc = Document()
    # 生成图表（保存到与docx同目录一致的临时目录）
    chart_dir = Path('/tmp/energy_report_charts')
    charts = build_report_charts(result_json, chart_dir)
    
    # ===== Cover =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('新能源电力解决方案分析报告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(10, 40, 100)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"\n{proj.get('project_name','')}\n\n")
    r2.font.size = Pt(14)
    r3 = p2.add_run(f"场景类型：{proj.get('scenario_type','')}    省份：{proj.get('province','')}")
    r3.font.size = Pt(10)
    
    # ===== 专家结论 =====
    doc.add_heading('专家分析结论', level=1)
    narrative = generate_narrative(result_json)
    for line in narrative.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    doc.add_page_break()
    
    # ===== 1. 项目概况 =====
    doc.add_heading('1. 项目概况', level=1)
    _add_table(doc, ['项目', '内容'], [
        ['项目名称', proj.get('project_name','')],
        ['场景类型', proj.get('scenario_type','')],
        ['省份/地区', proj.get('province','')],
        ['储能策略', d.get('storage_strategy_mode','')],
        ['数据完整度', f"{out.get('data_quality_results',{}).get('level','')} / {out.get('data_quality_results',{}).get('score','')}"],
    ])
    
    # ===== 2. 推荐方案 =====
    doc.add_heading('2. 推荐方案', level=1)
    _add_table(doc, ['设备', '规模', '说明'], [
        ['光伏', f"{s.get('pv_mwp') or 0} MWp", f"年发电 {float(sim.get('annual_pv_generation_mwh') or 0):,.0f} MWh"],
        ['风电', f"{s.get('wind_mw') or 0} MW", f"年发电 {float(sim.get('annual_wind_generation_mwh') or 0):,.0f} MWh"],
        ['储能', f"{s.get('storage_power_mw') or 0} MW / {s.get('storage_energy_mwh') or 0} MWh", f"年循环 {float(d.get('storage_equivalent_full_cycles_per_year') or 0):.0f} 次"],
    ])
    
    doc.add_paragraph(f"调度策略：{d.get('storage_strategy_mode','')}")
    doc.add_paragraph(f"日储能循环：{d.get('daily_storage_cycles')}    削峰量：{float(d.get('estimated_peak_reduction_kw') or 0):.0f} kW")
    
    # Monthly chart
    monthly = d.get('monthly_storage_revenue_breakdown', [])
    if monthly:
        doc.add_heading('储能月度充放电量', level=2)
        mrows = [[str(m['month']), f"{m['charge_mwh']:.0f}", f"{m['discharge_mwh']:.0f}"] for m in monthly]
        _add_table(doc, ['月份', '充电 (MWh)', '放电 (MWh)'], mrows)
        if charts.get('storage_monthly'):
            doc.add_picture(charts['storage_monthly'], width=Inches(6.5))
    
    # ===== 3. 发电量 =====
    doc.add_heading('3. 发电量及绿电覆盖率', level=1)
    pv_gen = float(sim.get('annual_pv_generation_mwh') or 0)
    wind_gen = float(sim.get('annual_wind_generation_mwh') or 0)
    gen_total = pv_gen + wind_gen
    cov = sim.get('coverage_ratio', 0)
    _add_table(doc, ['指标', '数值'], [
        ['年光伏发电量', f"{pv_gen:,.0f} MWh"],
        ['年风电发电量', f"{wind_gen:,.0f} MWh"],
        ['年总发电量', f"{gen_total:,.0f} MWh"],
        ['年储能放电量', f"{float(sim.get('annual_storage_discharge_mwh') or 0):,.0f} MWh"],
        ['年购电量', f"{float(sim.get('annual_grid_purchase_mwh') or 0):,.0f} MWh"],
        ['绿电覆盖率', f"{cov*100:.1f}%" if cov else 'N/A'],
        ['新能源自行消纳率', f"{sim.get('renewable_self_consumption_ratio',0)*100:.1f}%" if sim.get('renewable_self_consumption_ratio') else 'N/A'],
    ])
    
    # ===== 4. 财务分析 =====
    doc.add_heading('4. 财务分析', level=1)
    frows = [
        ['项目总投资 (CAPEX)', f"{f.get('capex_total',0)/1e4:,.0f} 万元"],
        ['年运营成本 (OPEX)', f"{f.get('opex_annual',0)/1e4:,.0f} 万元"],
        ['年收益/节费', f"{f.get('annual_savings_or_revenue',0)/1e4:,.0f} 万元"],
        ['Project IRR', f"{f['irr']*100:.2f}%" if f.get('irr') else 'N/A'],
        ['Equity IRR（70%负债）', f"{f.get('equity_irr',0)*100:.2f}%" if f.get('equity_irr') else 'N/A'],
        ['净现值 (NPV)', f"{f.get('npv',0)/1e4:,.0f} 万元"],
        ['权益净现值 (Equity NPV)', f"{f.get('equity_npv',0)/1e4:,.0f} 万元"],
        ['静态回收期', f"{f.get('payback_years','N/A')} 年"],
        ['动态回收期', f"{f.get('dyn_payback_years','N/A')} 年" if f.get('dyn_payback_years') else 'N/A'],
        ['总投资收益率 (ROI)', f"{f.get('roi')*100:.1f}%" if f.get('roi') is not None else 'N/A'],
        ['净资产收益率 (ROE)', f"{f.get('roe')*100:.1f}%" if f.get('roe') is not None else 'N/A'],
        ['偿债覆盖率 (DSCR)', f"{f.get('dscr_min','N/A')}"],
        ['系统平准化度电成本 (LCOE)', f"{f.get('lcoe','N/A')} 元/kWh"],
        ['平准化储能成本 (LCOS)', f"{f.get('lcos','N/A')} 元/kWh"],
        ['光伏 LCOE', f"{f.get('pv_lcoe','N/A')} 元/kWh"],
        ['风电 LCOE', f"{f.get('wind_lcoe','N/A')} 元/kWh"],
        ['储能 LCOS', f"{f.get('storage_lcos','N/A')} 元/kWh"],
    ]
    _add_table(doc, ['财务指标', '数值'], frows)
    if f.get('baseline_annual_energy_cost') is not None:
        doc.add_paragraph(f"基线年能源成本：{f.get('baseline_annual_energy_cost',0)/1e4:,.0f} 万元")
    if market.get('revenue_breakdown'):
        doc.add_heading('收益拆解', level=2)
        rb_rows = []
        for item in market.get('revenue_breakdown', []):
            if isinstance(item, dict):
                rb_rows.append([item.get('name',''), f"{item.get('amount',0)/1e4:,.0f} 万元"])
        if rb_rows:
            _add_table(doc, ['收益项', '金额'], rb_rows)
        if charts.get('revenue_breakdown'):
            doc.add_picture(charts['revenue_breakdown'], width=Inches(6.3))
        if charts.get('cost_structure'):
            doc.add_picture(charts['cost_structure'], width=Inches(5.6))
    if charts.get('cashflow'):
        doc.add_heading('项目现金流曲线', level=2)
        doc.add_picture(charts['cashflow'], width=Inches(6.5))
    
    # ===== 5. 融资分析 =====
    doc.add_heading('5. 融资结构分析', level=1)
    fin_mode = out.get('market_and_settlement', {}).get('financing_mode') or '贷款'
    doc.add_paragraph(f"融资模式：{fin_mode}")
    doc.add_paragraph(f"权益 IRR：{f.get('equity_irr',0)*100:.2f}%    （项目 IRR {f['irr']*100:.2f}%，杠杆放大 {f.get('equity_irr',0)/f['irr']:.1f}x）" if f.get('equity_irr') and f.get('irr') else "")
    doc.add_paragraph(f"DSCR 最低值：{f.get('dscr_min','N/A')}    （银行可贷标准：>1.2）")
    
    # ===== 6. 碳减排 =====
    doc.add_heading('6. 碳减排分析', level=1)
    _add_table(doc, ['指标', '数值'], [
        ['基线排放', f"{carbon.get('baseline_emissions_tco2e',0):,.0f} tCO₂e"],
        ['项目后排放', f"{carbon.get('post_project_emissions_tco2e',0):,.0f} tCO₂e"],
        ['年减排量', f"{carbon.get('annual_reduction_tco2e',0):,.0f} tCO₂e"],
        ['声明目标', carbon.get('claim_boundary_summary','') or '未设定'],
    ])
    
    # ===== 7. 风险分析 =====
    doc.add_heading('7. 风险与敏感性分析', level=1)
    
    # 单变量
    doc.add_heading('7.1 单变量敏感性', level=2)
    srows = []
    for item in sens:
        if 'impact_on_annual_revenue' in item:
            srows.append([item['factor'], f"{item['impact_on_annual_revenue']:,.0f} 元", item.get('irr_impact','')])
    if srows:
        _add_table(doc, ['因素', '收益影响', '敏感度'], srows)
    
    # 双变量热力图
    for item in sens:
        if 'heatmap' in item:
            doc.add_heading('7.2 双变量敏感性热力图（CAPEX × 发电量）', level=2)
            doc.add_paragraph('IRR 矩阵（%）:')
            hm = item['heatmap']
            headers = ['CAPEX\\发电量'] + [f"{g*100:.0f}%" for g in hm[0]['gen_factors']]
            hrows = []
            for row in hm:
                hrows.append([f"{row['capex_factor']*100:.0f}%"] + [f"{v*100:.1f}" if v else 'N/A' for v in row['irr_values']])
            _add_table(doc, headers, hrows)
    
    # 蒙特卡洛
    for item in sens:
        if 'monte_carlo' in item:
            doc.add_heading('7.3 蒙特卡洛 IRR 模拟', level=2)
            mc = item['monte_carlo']
            doc.add_paragraph(f"模拟次数：{mc.get('samples','500')} 次")
            _add_table(doc, ['P10（乐观）', 'P50（中位数）', 'P90（悲观）'], [
                [f"{mc['p10']*100:.1f}%", f"{mc['p50']*100:.1f}%", f"{mc['p90']*100:.1f}%"]
            ])
            doc.add_paragraph(f"置信度：90% 概率 IRR ≥ {mc['p10']*100:.1f}%")
    
    # 风险项
    if charts.get('tornado'):
        doc.add_heading('7.4 Tornado 敏感性图', level=2)
        doc.add_picture(charts['tornado'], width=Inches(6.5))
    if risks:
        doc.add_heading('7.5 风险提示', level=2)
        for r in risks:
            doc.add_paragraph(f"⚠ {r}", style='List Bullet')
    
    # ===== 8. 数据质量 =====
    doc.add_heading('8. 数据质量', level=1)
    dq = out.get('data_quality_results', {})
    doc.add_paragraph(f"数据质量评分：{dq.get('score','N/A')}/100    等级：{dq.get('level','N/A')}")
    for ck in dq.get('checks', []):
        doc.add_paragraph(f"[{ck['status']}] {ck['name']}: {ck['message']}", style='List Bullet')
    
    # ===== Footer =====
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— 报告完 —')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)
    
    docx_path = '/tmp/energy_report.docx'
    doc.save(docx_path)
    return docx_path

if __name__ == '__main__':
    import sys
    out = json.loads(open(sys.argv[1]).read())
    path = build_docx(out)
    print(f"✅ Word 报告已生成：{path}")
