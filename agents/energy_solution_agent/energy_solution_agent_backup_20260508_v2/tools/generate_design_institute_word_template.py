from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "templates" / "design_institute_report_template.docx"

TITLE_COLOR = RGBColor(31, 78, 121)
ACCENT_COLOR = RGBColor(47, 84, 150)
GRAY_COLOR = RGBColor(89, 89, 89)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for style_name, size, bold in [
        ("Title", 22, True),
        ("Heading 1", 15, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run("新能源电力解决方案分析报告")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = TITLE_COLOR
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(12)
    r2 = p2.add_run("（设计院版交付模板）")
    r2.font.size = Pt(14)
    r2.font.color.rgb = ACCENT_COLOR
    r2.font.name = "微软雅黑"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for text in [
        "项目名称：______________________________",
        "建设单位：______________________________",
        "编制单位：______________________________",
        "编制日期：______________________________",
        "适用场景：工商业储能 / 光储充 / 源网荷储 / 零碳工厂 / 微电网",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_page_break()


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], style: str = "Table Grid") -> None:
    doc.add_paragraph(title, style="Heading 3")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.name = "微软雅黑"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            for p in cell.paragraphs:
                for rr in p.runs:
                    rr.font.name = "微软雅黑"
                    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_paragraph("")


def add_chart_placeholder(doc: Document, title: str, note: str) -> None:
    doc.add_paragraph(title, style="Heading 3")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"【图表预留位】{title}")
    r.font.bold = True
    r.font.color.rgb = ACCENT_COLOR
    r.font.size = Pt(12)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = box.add_run("建议尺寸：宽 15~16 cm，高 7~9 cm\n（生成报告时由程序插入实际图表图片）")
    rr.font.color.rgb = GRAY_COLOR
    rr.font.name = "微软雅黑"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    n = doc.add_paragraph(f"说明：{note}")
    n.style = "Normal"
    doc.add_paragraph("")


def add_sections(doc: Document) -> None:
    doc.add_heading("1 项目概况与编制说明", level=1)
    doc.add_paragraph("1.1 项目背景：填写项目背景、边界、目标及建设必要性。")
    doc.add_paragraph("1.2 基础输入：项目名称、场景类型、省份/区域、分析模式、政策模式、电网区域、数据来源说明。")
    doc.add_paragraph("1.3 编制依据：政策文件、负荷数据、气象资源、设备参数、价格规则与假设。")

    add_table(doc, "表1-1 项目基础信息表", ["项目", "内容"], [
        ["项目名称", "${project_name}"],
        ["场景类型", "${scenario_type}"],
        ["场景能力分类", "${scenario_detail_label}"],
        ["运行玩法", "${operation_mode}"],
        ["分析模式", "${analysis_mode}"],
        ["省份/区域", "${province}"],
        ["数据完整度", "${data_completeness_grade}"],
        ["数据质量等级", "${data_quality_level} / ${data_quality_score}"],
        ["省级规则状态", "${province_profile_status}"],
        ["政策模式", "${market_policy_mode}"],
        ["电网区域", "${profile_grid_region}"],
    ])

    doc.add_heading("2 资源条件与负荷特性分析", level=1)
    add_table(doc, "表2-1 资源条件汇总表", ["指标", "数值"], [
        ["光伏资源口径", "${pv_resource_basis} / ${pv_resource_accuracy}"],
        ["光伏 P50/P90", "${pv_p50_generation_mwh} / ${pv_p90_generation_mwh}"],
        ["光伏有效PR", "${pv_pr_effective}"],
        ["光伏LCOE", "${pv_lcoe_per_kwh}"],
        ["风电资源口径", "${wind_resource_basis} / ${wind_resource_accuracy}"],
        ["风电 P50/P90", "${wind_p50_generation_mwh} / ${wind_p90_generation_mwh}"],
        ["风电功率曲线", "${wind_power_curve_used}"],
        ["风电LCOE", "${wind_lcoe_per_kwh}"],
    ])
    add_chart_placeholder(doc, "图2-1 年度负荷曲线及典型日负荷曲线", "展示原始负荷、净负荷、典型工作日/周末曲线。")
    add_chart_placeholder(doc, "图2-2 光伏/风电资源出力特性图", "展示月度或典型日资源特性、P50/P90对比。")

    doc.add_heading("3 方案配置与技术路线", level=1)
    add_table(doc, "表3-1 推荐方案配置表", ["项目", "推荐值"], [
        ["光伏", "${pv_mwp} MWp"],
        ["风电", "${wind_mw} MW"],
        ["储能", "${storage_power_mw} MW / ${storage_energy_mwh} MWh"],
        ["储能实际值", "${raw_storage_power_mw} MW / ${raw_storage_energy_mwh} MWh"],
        ["充电容量", "${charging_capacity_mw} MW"],
        ["供冷容量", "${cooling_capacity_rt} RT"],
        ["供热容量", "${heating_capacity_mwth} MWth"],
    ])
    add_chart_placeholder(doc, "图3-1 系统构成与能量流向示意图", "可由外部绘图或流程图输出后插入。")

    doc.add_heading("4 运行策略与调度分析", level=1)
    add_table(doc, "表4-1 调度结果汇总表", ["指标", "结果"], [
        ["基线峰值购电功率", "${baseline_peak_grid_kw}"],
        ["储能后峰值购电功率", "${post_storage_peak_grid_kw}"],
        ["估算削峰量", "${estimated_peak_reduction_kw}"],
        ["日储能循环次数", "${daily_storage_cycles}"],
        ["储能策略模式", "${storage_strategy_mode}"],
        ["储能容量测算口径", "${storage_sizing_basis}"],
        ["年储能吞吐量", "${storage_annual_throughput_mwh}"],
        ["年等效满循环", "${storage_equivalent_full_cycles_per_year}"],
        ["估算寿命年限", "${storage_life_years_estimate}"],
        ["有效往返效率", "${storage_effective_round_trip_efficiency}"],
    ])
    add_chart_placeholder(doc, "图4-1 典型日功率平衡图", "建议包含负荷、光伏、风电、储能充放电、电网购电。")
    add_chart_placeholder(doc, "图4-2 储能SOC变化图", "建议按典型日或代表周展示SOC、充放电窗口与循环次数。")
    add_chart_placeholder(doc, "图4-3 月度收益/吞吐量图", "展示12个月储能收益、充放电量、毛收益。")

    doc.add_heading("5 电力市场与结算分析", level=1)
    add_table(doc, "表5-1 市场关键参数表", ["参数", "取值"], [
        ["省级规则状态", "${province_profile_status}"],
        ["市场政策模式", "${market_policy_mode}"],
        ["交易价差", "${trading_price_spread_per_kwh}"],
        ["价格波动指数", "${trading_volatility_index}"],
        ["最优月份/最弱月份", "${trading_best_month} / ${trading_worst_month}"],
        ["在线规则生效字段", "${live_rule_effective_fields}"],
    ])
    add_chart_placeholder(doc, "图5-1 分时电价/现货价格曲线图", "展示峰平谷或实时电价曲线。")
    add_chart_placeholder(doc, "图5-2 套利窗口识别图", "展示充电窗口、放电窗口、价差阈值与执行结果。")

    doc.add_heading("6 财务测算与经济性评价", level=1)
    add_table(doc, "表6-1 财务测算汇总表", ["指标", "结果"], [
        ["年收益/节费", "${annual_savings_or_revenue}"],
        ["年税费合计", "${annual_tax_total}"],
        ["年企业所得税", "${annual_income_tax}"],
        ["年增值税（应纳）", "${annual_vat_payable}"],
        ["年附加税", "${annual_vat_surcharges_only}"],
        ["年电量电费", "${annual_energy_charge_cost}"],
        ["年需量电费", "${annual_demand_charge_cost}"],
        ["年辅助服务收益", "${annual_ancillary_service_revenue}"],
        ["年需求响应收益", "${annual_demand_response_revenue}"],
        ["年外送收益", "${annual_export_revenue}"],
        ["项目IRR", "${irr}"],
        ["回收期", "${payback_years}"],
        ["NPV", "${npv}"],
        ["单位减碳成本", "${abatement_cost_per_tco2e}"],
    ])
    add_chart_placeholder(doc, "图6-1 项目全生命周期现金流图", "展示年度现金流、累计现金流、回收期位置。")
    add_chart_placeholder(doc, "图6-2 收益构成瀑布图/堆叠图", "展示削峰、套利、辅助服务、需求响应等收益构成。")
    add_chart_placeholder(doc, "图6-3 敏感性分析龙卷风图", "展示关键因素对IRR/NPV的影响。")

    doc.add_heading("7 碳减排效益分析", level=1)
    add_table(doc, "表7-1 碳结果汇总表", ["指标", "结果"], [
        ["基线排放", "${baseline_emissions_tco2e}"],
        ["项目后排放", "${post_project_emissions_tco2e}"],
        ["年减排量", "${annual_reduction_tco2e}"],
        ["声明边界", "${claim_boundary_summary}"],
    ])
    add_chart_placeholder(doc, "图7-1 碳排放对比图", "展示基线与项目后排放对比。")
    add_chart_placeholder(doc, "图7-2 减排路径拆分图", "展示光伏、风电、储能、热泵等路径减排贡献。")

    doc.add_heading("8 风险分析与实施建议", level=1)
    doc.add_paragraph("8.1 风险清单：填写政策、电价、负荷、资源、设备、施工、并网等风险。")
    doc.add_paragraph("8.2 数据缺口：填写需补充的实测数据、协议、图纸与规则文件。")
    doc.add_paragraph("8.3 实施建议：分近期、中期、远期给出建议。")
    add_table(doc, "表8-1 风险与应对措施表", ["风险项", "影响", "建议措施"], [
        ["${risk_1}", "${risk_impact_1}", "${risk_action_1}"],
        ["${risk_2}", "${risk_impact_2}", "${risk_action_2}"],
        ["${risk_3}", "${risk_impact_3}", "${risk_action_3}"],
    ])

    doc.add_heading("9 结论与建议", level=1)
    doc.add_paragraph("9.1 结论摘要：填写建议建设规模、收益水平、减排效益与可实施性。")
    doc.add_paragraph("9.2 下一步工作建议：可研深化、现场踏勘、并网咨询、设备选型、财务尽调等。")

    doc.add_heading("附录A 主要输入参数表", level=1)
    add_table(doc, "附表A-1 输入参数总表", ["分类", "参数", "值", "备注"], [
        ["项目", "project_name", "${project_name}", ""],
        ["市场", "market_policy_mode", "${market_policy_mode}", ""],
        ["调度", "storage_strategy_mode", "${storage_strategy_mode}", ""],
        ["资源", "pv_resource_basis", "${pv_resource_basis}", ""],
        ["财务", "irr", "${irr}", "输出结果示例"],
    ])

    doc.add_heading("附录B 图表清单（程序自动插入）", level=1)
    for item in [
        "图2-1 年度负荷曲线及典型日负荷曲线",
        "图2-2 光伏/风电资源出力特性图",
        "图3-1 系统构成与能量流向示意图",
        "图4-1 典型日功率平衡图",
        "图4-2 储能SOC变化图",
        "图4-3 月度收益/吞吐量图",
        "图5-1 分时电价/现货价格曲线图",
        "图5-2 套利窗口识别图",
        "图6-1 项目全生命周期现金流图",
        "图6-2 收益构成瀑布图/堆叠图",
        "图6-3 敏感性分析龙卷风图",
        "图7-1 碳排放对比图",
        "图7-2 减排路径拆分图",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_page_number(doc.sections[0])
    add_cover(doc)
    add_toc(doc)
    add_sections(doc)
    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
