#!/usr/bin/env python3
"""
储能企业竞品调研报告生成器
生成专业的Word格式竞品分析报告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_chinese_font(run, font_name='Microsoft YaHei', size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_chinese(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    if level == 1:
        set_chinese_font(run, 'Microsoft YaHei', 18, True)
        run.font.color.rgb = RGBColor(26, 58, 108)
    elif level == 2:
        set_chinese_font(run, 'Microsoft YaHei', 14, True)
        run.font.color.rgb = RGBColor(44, 82, 130)
    else:
        set_chinese_font(run, 'Microsoft YaHei', 12, True)
    
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_chinese(doc, text, bold=False, size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, 'Microsoft YaHei', size, bold)
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    return p

def create_competitor_report():
    """创建竞品调研报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    run = title.add_run('中国储能行业\n竞品调研报告')
    set_chinese_font(run, 'Microsoft YaHei', 28, True)
    run.font.color.rgb = RGBColor(26, 58, 108)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('—— 主流储能企业产品对比分析')
    set_chinese_font(run, 'Microsoft YaHei', 16, False)
    run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    run = date_p.add_run(f'报告日期：2026年3月')
    set_chinese_font(run, 'Microsoft YaHei', 12, False)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分页
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_heading_chinese(doc, '目录', 1)
    
    toc_items = [
        '一、调研背景与目的',
        '二、储能行业概况',
        '三、主流企业产品对比',
        '四、技术参数详细对比',
        '五、价格与商业模式分析',
        '六、竞争优势分析',
        '七、市场机会与建议',
    ]
    
    for item in toc_items:
        add_paragraph_chinese(doc, item)
    
    doc.add_page_break()
    
    # ========== 一、调研背景与目的 ==========
    add_heading_chinese(doc, '一、调研背景与目的', 1)
    
    add_paragraph_chinese(doc, '1.1 调研背景')
    add_paragraph_chinese(doc, 
        '随着全球能源转型加速，储能产业进入快速发展期。2025年中国储能装机规模突破100GW，'
        '工商业储能和独立储能成为主要增长点。各大储能企业纷纷推出新品，技术路线趋于多元化，'
        '市场竞争日趋激烈。')
    
    add_paragraph_chinese(doc, '1.2 调研目的')
    purposes = [
        '• 梳理主流储能企业的最新产品布局',
        '• 对比分析各企业产品的技术参数和性能指标',
        '• 分析不同企业的定价策略和商业模式',
        '• 为企业储能项目选型提供决策参考',
    ]
    for p in purposes:
        add_paragraph_chinese(doc, p)
    
    # ========== 二、储能行业概况 ==========
    add_heading_chinese(doc, '二、储能行业概况', 1)
    
    add_paragraph_chinese(doc, '2.1 市场规模')
    add_paragraph_chinese(doc,
        '2025年中国储能市场规模超过3000亿元，其中电化学储能占比超过90%。'
        '工商业储能、独立储能、光储充一体化成为三大主流应用场景。')
    
    add_paragraph_chinese(doc, '2.2 技术路线')
    tech_routes = [
        '• 磷酸铁锂（LFP）：主流技术路线，占比超95%，安全性高、成本低',
        '• 钠离子电池：新兴技术，原材料成本低，适合大规模储能',
        '• 液流电池：长时储能首选，循环寿命长，但初投资高',
        '• 压缩空气/飞轮：物理储能，适用于特定场景',
    ]
    for t in tech_routes:
        add_paragraph_chinese(doc, t)
    
    doc.add_page_break()
    
    # ========== 三、主流企业产品对比 ==========
    add_heading_chinese(doc, '三、主流企业产品对比', 1)
    
    # 创建企业对比表
    add_heading_chinese(doc, '3.1 企业概览', 2)
    
    companies = [
        ['企业名称', '股票代码', '总部', '主营业务', '2024年储能营收', '市场地位'],
        ['宁德时代', '300750', '福建宁德', '动力电池、储能电池', '~800亿元', '全球第一'],
        ['比亚迪', '002594', '广东深圳', '新能源汽车、储能', '~600亿元', '全球第二'],
        ['亿纬锂能', '300014', '广东惠州', '锂电池、储能系统', '~200亿元', '国内前三'],
        ['海辰储能', '未上市', '福建厦门', '储能电池、储能系统', '~150亿元', '储能新秀'],
        ['瑞浦兰钧', '00666.HK', '浙江温州', '动力电池、储能电池', '~120亿元', '快速崛起'],
        ['国轩高科', '002074', '安徽合肥', '动力电池、储能', '~100亿元', '老牌企业'],
        ['鹏辉能源', '300438', '广东广州', '锂电池、储能', '~80亿元', '细分领域'],
        ['南都电源', '300068', '浙江杭州', '储能系统、铅酸电池', '~60亿元', '系统集成'],
    ]
    
    table = doc.add_table(rows=len(companies), cols=6)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(companies):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'Microsoft YaHei', 9, i==0)
            if i == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2C5282')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # ========== 四、技术参数详细对比 ==========
    add_heading_chinese(doc, '四、技术参数详细对比', 1)
    
    add_heading_chinese(doc, '4.1 电芯产品对比', 2)
    
    cells = [
        ['企业', '产品型号', '容量(Ah)', '能量密度(Wh/kg)', '循环寿命', '技术特点'],
        ['宁德时代', '280Ah LFP', '280', '160-170', '8000+', '高安全、长寿命'],
        ['宁德时代', '306Ah LFP', '306', '165-175', '10000+', '大容量、高集成'],
        ['宁德时代', '314Ah LFP', '314', '170-180', '12000+', '最新一代'],
        ['比亚迪', '刀片电池', '100+', '140-150', '8000+', 'CTP技术、高安全'],
        ['亿纬锂能', 'LF280K', '280', '155-165', '8000+', '性价比高'],
        ['亿纬锂能', 'LF560K', '560', '160-170', '10000+', '大容量新品'],
        ['海辰储能', '300Ah', '300', '160-170', '10000+', '储能专用'],
        ['瑞浦兰钧', '问顶320Ah', '320', '165-175', '10000+', '高能量密度'],
    ]
    
    table = doc.add_table(rows=len(cells), cols=6)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(cells):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'Microsoft YaHei', 9, i==0)
            if i == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2C5282')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    add_heading_chinese(doc, '4.2 储能系统产品对比', 2)
    
    systems = [
        ['企业', '系统产品', '容量(MWh)', '应用场景', '关键特性', '交付方式'],
        ['宁德时代', 'EnerOne', '2.5-5', '工商业', '模块化、高安全', '集装箱'],
        ['宁德时代', 'EnerC', '10-20', '大型储能', '高集成、长寿命', '预制舱'],
        ['比亚迪', 'Cube Pro', '2.5-5', '工商业', '刀片电池、液冷', '集装箱'],
        ['比亚迪', 'MC-I', '10+', '大型储能', '高能量密度', '预制舱'],
        ['亿纬锂能', 'LF P50', '2.5', '工商业', '标准化、易扩展', '户外柜'],
        ['海辰储能', 'HiTHIUM', '2.5-5', '工商业/大型', '储能专用电芯', '集装箱'],
        ['瑞浦兰钧', '储能集装箱', '2.5-20', '全场景', '问顶技术', '定制'],
    ]
    
    table = doc.add_table(rows=len(systems), cols=6)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(systems):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'Microsoft YaHei', 9, i==0)
            if i == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2C5282')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_page_break()
    
    # ========== 五、价格与商业模式分析 ==========
    add_heading_chinese(doc, '五、价格与商业模式分析', 1)
    
    add_heading_chinese(doc, '5.1 产品价格区间', 2)
    
    prices = [
        ['产品类型', '价格区间(元/Wh)', '代表企业', '价格趋势'],
        ['储能电芯(LFP)', '0.35-0.50', '宁德时代、比亚迪', '持续下降'],
        ['工商业储能系统', '0.80-1.20', '宁德时代、海辰', '逐步下降'],
        ['大型储能系统', '0.60-0.90', '宁德时代、比亚迪', '竞争加剧'],
        ['储能集成(EPC)', '1.00-1.50', '各类集成商', '略有下降'],
    ]
    
    table = doc.add_table(rows=len(prices), cols=4)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(prices):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'Microsoft YaHei', 9, i==0)
            if i == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2C5282')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    add_heading_chinese(doc, '5.2 商业模式对比', 2)
    
    models = [
        '• 宁德时代："电芯+系统"双轮驱动，提供整体解决方案',
        '• 比亚迪：垂直整合，从电芯到系统集成全链条覆盖',
        '• 亿纬锂能：聚焦电芯和模组，与系统集成商合作',
        '• 海辰储能：专注储能领域，提供高性价比储能专用产品',
        '• 瑞浦兰钧：技术驱动，问顶技术提升产品竞争力',
    ]
    for m in models:
        add_paragraph_chinese(doc, m)
    
    # ========== 六、竞争优势分析 ==========
    add_heading_chinese(doc, '六、竞争优势分析', 1)
    
    add_heading_chinese(doc, '6.1 各企业核心竞争力', 2)
    
    advantages = [
        ['企业', '核心优势', '技术亮点', '主要客户', '市场策略'],
        ['宁德时代', '规模+技术', '高安全电芯、长寿命', '国家电网、五大发电', '高端市场'],
        ['比亚迪', '垂直整合', '刀片电池、CTP', '自用+外部客户', '全产业链'],
        ['亿纬锂能', '成本控制', '性价比产品', '中小集成商', '性价比路线'],
        ['海辰储能', '储能专注', '储能专用电芯', '储能集成商', '细分市场'],
        ['瑞浦兰钧', '技术创新', '问顶技术、高能量', '新兴客户', '技术突破'],
    ]
    
    table = doc.add_table(rows=len(advantages), cols=5)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(advantages):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, 'Microsoft YaHei', 9, i==0)
            if i == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '2C5282')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    doc.add_page_break()
    
    # ========== 七、市场机会与建议 ==========
    add_heading_chinese(doc, '七、市场机会与建议', 1)
    
    add_heading_chinese(doc, '7.1 选型建议', 2)
    
    recommendations = [
        '【高端项目】选择宁德时代或比亚迪',
        '  - 适用：对安全性、可靠性要求极高的项目',
        '  - 优势：技术领先、品牌背书、售后服务完善',
        '  - 注意：价格相对较高，需要评估性价比',
        '',
        '【性价比项目】选择亿纬锂能或海辰储能',
        '  - 适用：预算有限但要求性能稳定的项目',
        '  - 优势：价格适中，性能满足主流需求',
        '  - 注意：品牌影响力相对较弱',
        '',
        '【技术创新项目】关注瑞浦兰钧',
        '  - 适用：愿意尝试新技术、追求高能量密度的项目',
        '  - 优势：技术突破、产品创新',
        '  - 注意：市场验证时间较短',
    ]
    
    for r in recommendations:
        if r:
            add_paragraph_chinese(doc, r, '【' in r)
    
    add_heading_chinese(doc, '7.2 市场趋势判断', 2)
    
    trends = [
        '• 电芯容量持续增大：280Ah→300Ah→500Ah+，降低系统集成成本',
        '• 系统集成度提升：预制舱、一体化设计成为主流',
        '• 智能化水平提高：BMS、EMS、AI运维深度融合',
        '• 价格持续下降：规模化效应+技术进步，预计每年下降10-15%',
        '• 长时储能需求增长：4小时以上储能系统需求增加',
    ]
    for t in trends:
        add_paragraph_chinese(doc, t)
    
    # 结语
    doc.add_paragraph()
    add_heading_chinese(doc, '结语', 1)
    
    add_paragraph_chinese(doc,
        '储能行业正处于快速发展期，技术迭代加速，市场竞争激烈。'
        '建议企业在选型时综合考虑技术性能、价格成本、品牌服务等因素，'
        '选择最适合项目需求的合作伙伴。同时关注行业技术发展趋势，'
        '为未来项目升级预留空间。')
    
    # 保存
    output_path = '/Users/zhaoruicn/.openclaw/workspace/储能企业竞品调研报告_2026年3月.docx'
    doc.save(output_path)
    
    print(f"✅ 竞品调研报告已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    create_competitor_report()
