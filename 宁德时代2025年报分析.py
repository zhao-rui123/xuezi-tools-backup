from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run, font_name='SimSun', size=10.5, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

# 标题
title = doc.add_heading('宁德时代新能源科技股份有限公司', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_chinese_font(title.runs[0], 'SimHei', 18, True)

subtitle = doc.add_heading('2025年年度报告深度分析', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_chinese_font(subtitle.runs[0], 'SimHei', 16, True)

# 分析日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('分析日期：2026年3月20日')
set_chinese_font(date_run, 'SimSun', 10.5)

doc.add_paragraph()  # 空行

# 一、公司概况
doc.add_heading('一、公司概况与行业地位', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
宁德时代（CATL）是全球领先的新能源创新科技公司，专注于动力电池和储能电池的研发、生产和销售。

【核心地位】
• 动力电池出货量连续全球第一
• 储能电池出货量全球领先
• 2025年锂电池销量达661GWh，同比增长近40%
• 连续九个季度入选BNEF Tier 1名单
• 标普全球2025年榜单位列全球第八位

【港股IPO】
2025年成功完成港股IPO，树立全球资本市场标杆项目，为海外战略布局提供稳健资本支持。
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 二、财务分析
doc.add_heading('二、财务业绩分析', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

# 添加表格
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'

# 表头
headers = ['指标', '2025年数据', '同比变化']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_chinese_font(cell.paragraphs[0].runs[0], 'SimHei', 10.5, True)

# 数据行
data = [
    ['锂电池销量', '661 GWh', '增长近40%'],
    ['全球储能累计交付量', '突破20 GWh', '持续增长'],
    ['废旧电池回收量', '21万吨', '增长超60%'],
    ['换电站建设', '1,325座', '快速扩张'],
    ['累计现金分红', '近千亿元', '连续三年50%净利润分红']
]

for i, row_data in enumerate(data, 1):
    for j, text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = text
        set_chinese_font(cell.paragraphs[0].runs[0], 'SimSun', 10.5)

doc.add_paragraph()

# 三、业务板块分析
doc.add_heading('三、核心业务板块分析', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

## 3.1 动力电池业务
doc.add_heading('3.1 动力电池业务', level=2)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 12, True)

content = """
【市场地位】
• 全球动力电池出货量第一
• 客户覆盖全球主流车企
• 技术授权、联合研发、本地建厂多元化合作模式

【新产品发布】
• "钠新"电池：优异低温表现，降低锂资源依赖
• "骁遥"双核系列：突破单一化学体系性能局限
• "天行"商用车电池：解决续航短、补能慢、衰减快痛点
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

## 3.2 储能电池业务
doc.add_heading('3.2 储能电池业务', level=2)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 12, True)

content = """
【市场地位】
• 2025年上半年全球储能电芯出货量第一（宁德时代）
• 进入英国、德国、意大利等欧洲主流市场出货量前十
• 截至2025年底全球储能累计交付量突破20GWh

【重点项目】
• 阿联酋19GWh项目
• 澳大利亚24GWh长时储能项目
• 587Ah大电芯、TENER Stack 9MWh储能系统

【技术突破】
• "天恒""TENER Stack"大容量储能系统：实现安全寿命最优解
• 587Ah储能专用电芯开启量产交付
• 500+Ah储能大电芯量产能力
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

## 3.3 换电业务
doc.add_heading('3.3 换电业务布局', level=2)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 12, True)

content = """
【换电站建设】
• 截至2025年底累计建成换电站：1,325座
  - 巧克力换电站：突破1,000座
  - 骐骥换电站：突破300座

【合作伙伴】
• 乘用车：广汽、长安、一汽、上汽、奇瑞
• 商用车：一汽解放、陕重汽
• 能源伙伴：中石化

【应用场景】
• 乘用车便捷补能
• 商用车干线物流
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 四、技术创新
doc.add_heading('四、技术创新与研发实力', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【研发投入】
• 研发人员：23,000余名
• 国际专利申请量：中国企业第二
• 宁德市全球创新强度排名：第四位

【AI融合创新】
• 荣获世界经济论坛（WEF）"AI驱动产业转型全球标杆"MINDS奖项
• AI深度融合到创新和制造过程

【灯塔工厂】
• WEF评选行业最多的灯塔工厂
• 唯一"可持续灯塔工厂"
• 关键质量标准较行业平均实现数量级领先

【前沿技术储备】
• 零碳电网技术
• 电力电子技术
• 柔性调控技术
• 虚拟电厂技术
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 五、新兴领域拓展
doc.add_heading('五、新兴领域拓展', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【电动船舶】
• 搭载宁德时代电池的船舶近900艘
• 纯电船舶将驶向远洋

【电动航空】
• 旗下峰飞航空2吨级eVTOL完成多次复杂环境飞行验证
• 全球最大5吨级eVTOL完成首次公开飞行

【新能源重卡】
• "天行"电池成为主流选择
• 与多家重卡企业联合推出换电车型

【零碳产业园】
• 山东东营、江苏盐城、福建宁德、四川宜宾等地推动建设
• 助力钢铁、水泥、化工等传统行业转型升级
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 六、ESG与可持续发展
doc.add_heading('六、ESG表现与可持续发展', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【碳中和目标】
• 2025年：核心运营电池工厂均实现碳中和 ✓
• 2035年：价值链碳中和（稳步迈进）

【ESG评级】
• MSCI评级：AA级（行业领先）
• 首次入选标普全球《可持续发展年鉴》

【循环经济】
• 旗下邦普循环2025年回收废旧电池及材料21万吨
• 同比增长超60%

【社会责任】
• 携峰飞航空向香港大埔火灾救援捐赠1,500万港元
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 七、全球化战略
doc.add_heading('七、全球化战略布局', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【资本运作】
• 港股IPO成功，树立全球资本市场标杆
• 与全球投资者共享公司成长

【本地化运营】
• 技术授权、联合研发、本地建厂
• 为客户提供定制化产品和本地化服务

【海外产能】
• 海辰储能美国得州工厂：年产10GWh（2025年7月投产）
• 远景动力美国田纳西州工厂：一期7GWh（已投产）
• 中创新航葡萄牙基地：首期15GWh（2025年2月动工）
• 远景动力西班牙超级工厂：2026年投产，欧洲首个磷酸铁锂电池超级工厂

【市场拓展】
• 欧洲：英国、德国、意大利等主流市场
• 澳洲：24GWh长时储能项目
• 中东：阿联酋19GWh项目
• 日本：户用储能市场加速拓展
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 八、投资亮点
doc.add_heading('八、投资亮点与核心优势', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【核心竞争优势】
1. 规模优势：全球动力电池+储能电池双龙头
2. 技术优势：23,000+研发人员，国际专利中国企业第二
3. 质量优势：关键质量标准行业领先，最多灯塔工厂
4. 成本优势：高度复杂的系统工程能力，持续降低系统成本
5. 客户优势：全球主流车企和储能客户全覆盖

【增长驱动因素】
1. 全球新能源渗透率持续提升
2. 储能市场高速增长（2025年全球户储出货量同比+50%）
3. 换电业务模式创新
4. 新兴领域（船舶、航空、重卡）拓展
5. 全球化产能布局加速

【股东回报】
• 连续三年以净利润50%实施现金分红
• 累计分红接近千亿元
• 高强度科技创新+现金分红双轮驱动
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 九、风险提示
doc.add_heading('九、风险提示', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【行业风险】
• 新能源产业周期性波动
• 原材料价格（锂、镍、钴）波动
• 技术路线迭代风险（固态电池、钠离子电池等）

【竞争风险】
• 行业竞争加剧，价格战压力
• 比亚迪、中创新航、海辰储能等竞争对手追赶
• 韩系企业（LGES、三星SDI）市场份额恢复

【地缘政治风险】
• 国际贸易摩擦
• 海外产能建设不确定性
• 供应链本土化要求

【政策风险】
• 各国新能源补贴政策变化
• 碳关税等贸易壁垒
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 十、总结与展望
doc.add_heading('十、总结与展望', level=1)
set_chinese_font(doc.paragraphs[-1].runs[0], 'SimHei', 14, True)

content = """
【2025年业绩总结】
宁德时代2025年交出了一份亮眼的成绩单：
• 锂电池销量661GWh，同比增长近40%
• 动力+储能双轮驱动，全球龙头地位稳固
• 港股IPO成功，资本运作能力凸显
• 核心运营实现碳中和，ESG表现行业领先
• 累计分红近千亿元，股东回报丰厚

【未来展望】
根据国际能源署（IEA）预测，到2050年实现净零排放，全球年度能源投资将达4.5万亿美元。

宁德时代正从"局部突破"走向"全域增量"：
• 产品从乘用车延伸至商用车、船舶、航空
• 从电池产品向能源服务演进（换电、零碳园区）
• 全球化产能布局加速，海外市场份额提升
• 技术创新持续突破（大电芯、钠电池、双核系列）

【投资结论】
作为全球新能源产业的领军企业，宁德时代具备：
✓ 强大的技术护城河
✓ 规模化的成本优势
✓ 全球化的客户基础
✓ 持续的创新能力
✓ 稳健的股东回报

长期投资价值显著，建议持续关注。
"""
p = doc.add_paragraph(content.strip())
set_chinese_font(p.runs[0], 'SimSun', 10.5)

# 页脚
doc.add_paragraph()
doc.add_paragraph('— 报告完 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
doc.save('/Users/zhaoruicn/.openclaw/workspace/宁德时代2025年报分析报告.docx')
print("报告已生成：/Users/zhaoruicn/.openclaw/workspace/宁德时代2025年报分析报告.docx")
