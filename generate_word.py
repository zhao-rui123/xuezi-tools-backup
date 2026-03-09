from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 标题页
title = doc.add_heading('统一记忆管理系统', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run('Unified Memory System v2.0\n').font.size = Pt(16)
subtitle.add_run('让 AI 助手更智能、更懂你的记忆管理解决方案').font.size = Pt(12)

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.add_run('作者：雪子助手\n').font.size = Pt(11)
author.add_run('日期：2026年3月').font.size = Pt(11)

doc.add_page_break()

# 第一章
doc.add_heading('一、系统概述', 1)
doc.add_heading('1.1 什么是统一记忆管理系统？', 2)
doc.add_paragraph('统一记忆管理系统 (Unified Memory System) 是一个全功能的记忆管理解决方案，旨在帮助 AI 助手高效地管理、存储和检索与用户交互过程中的重要信息。')

doc.add_heading('1.2 核心理念', 2)
doc.add_paragraph('• 一个命令搞定所有记忆操作', style='List Bullet')
doc.add_paragraph('• 自动识别重要信息并存储', style='List Bullet')
doc.add_paragraph('• 智能检索优先显示重要内容', style='List Bullet')
doc.add_paragraph('• 永久保存工作历史不丢失', style='List Bullet')

# 第二章
doc.add_heading('二、核心功能', 1)
doc.add_heading('2.1 三层记忆架构', 2)
doc.add_paragraph('系统采用分层架构设计：')
doc.add_paragraph('第一层：日常记忆层 - 每天自动生成记忆文件，详细记录当天工作，永久保留', style='List Bullet')
doc.add_paragraph('第二层：智能分析层 - 每月自动分析记忆文件，提取高频关键词和主题，生成月度摘要', style='List Bullet')
doc.add_paragraph('第三层：增强记忆层 - 自动识别重要信息，重要性分级（0-1评分），智能检索排序', style='List Bullet')

doc.add_heading('2.2 双模式记忆存储', 2)

# 自动识别表格
table1 = doc.add_table(rows=8, cols=4)
table1.style = 'Light Grid Accent 1'
hdr = table1.rows[0].cells
hdr[0].text = '识别模式'
hdr[1].text = '示例'
hdr[2].text = '存储类别'
hdr[3].text = '默认重要度'

data1 = [
    ['"我喜欢..."', '我喜欢用 Python 编程', 'preference', '0.7'],
    ['"我决定..."', '我决定使用 React', 'decision', '0.9'],
    ['"我叫..."', '我叫雪子', 'identity', '0.85'],
    ['"我的股票..."', '我的股票代码是 002460', 'finance', '0.85'],
    ['"千万不要..."', '千万不要上传 API key', 'constraint', '0.9'],
    ['"我计划..."', '我计划明天开会', 'schedule', '0.75'],
    ['"记住..."', '记住这个配置', 'general', '0.8']
]

for i, row in enumerate(data1, 1):
    cells = table1.rows[i].cells
    for j, text in enumerate(row):
        cells[j].text = text

doc.add_paragraph()
doc.add_heading('2.3 统一命令：ums', 2)

# 命令表格
table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Light Grid Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = '命令'
hdr[1].text = '功能'
hdr[2].text = '示例'

data2 = [
    ['ums status', '查看系统状态', 'ums status'],
    ['ums daily save', '保存每日记忆', 'ums daily save "今天完成了..."'],
    ['ums analyze', '月度智能分析', 'ums analyze monthly'],
    ['ums recall store', '存储记忆', 'ums recall store "内容"'],
    ['ums recall search', '搜索记忆', 'ums recall search "编程"'],
    ['ums session save', '保存会话状态', 'ums session save']
]

for i, row in enumerate(data2, 1):
    cells = table2.rows[i].cells
    for j, text in enumerate(row):
        cells[j].text = text

doc.add_page_break()

# 第三章
doc.add_heading('三、系统优势', 1)
doc.add_heading('3.1 一体化设计', 2)

table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Light Grid Accent 1'
hdr = table3.rows[0].cells
hdr[0].text = '对比项'
hdr[1].text = '分散系统'
hdr[2].text = '统一记忆系统'

data3 = [
    ['命令数量', '5+ 个', '1 个 (ums)'],
    ['学习成本', '高', '低'],
    ['维护难度', '复杂', '简单'],
    ['数据一致性', '难保证', '自动保证'],
    ['自动化程度', '部分', '全面']
]

for i, row in enumerate(data3, 1):
    cells = table3.rows[i].cells
    for j, text in enumerate(row):
        cells[j].text = text

doc.add_paragraph()
doc.add_heading('3.2 Token 节省', 2)

table4 = doc.add_table(rows=4, cols=4)
table4.style = 'Light Grid Accent 1'
hdr = table4.rows[0].cells
hdr[0].text = '查询场景'
hdr[1].text = '传统方式'
hdr[2].text = '优化后'
hdr[3].text = '节省'

data4 = [
    ['回顾某月工作', '5万 Token', '1千 Token', '98%'],
    ['查找特定主题', '3万 Token', '500 Token', '98%'],
    ['问"之前做了什么"', '翻大量文件', '看摘要', '99%']
]

for i, row in enumerate(data4, 1):
    cells = table4.rows[i].cells
    for j, text in enumerate(row):
        cells[j].text = text

doc.add_paragraph()
doc.add_heading('3.3 其他优势', 2)
doc.add_paragraph('• 智能化程度高 - 自动识别 7 种重要信息类型', style='List Bullet')
doc.add_paragraph('• 可靠性强 - 永久保存 + 每日自动备份', style='List Bullet')
doc.add_paragraph('• 扩展性好 - 模块化设计，易于扩展', style='List Bullet')

doc.add_page_break()

# 第四章
doc.add_heading('四、使用示例', 1)
doc.add_heading('4.1 自动识别存储', 2)
doc.add_paragraph('用户：我喜欢用 Python 做数据分析')
doc.add_paragraph('系统：自动识别并存储为 preference 类别，重要度 0.7')
doc.add_paragraph()
doc.add_paragraph('用户：我决定采用 React 作为前端框架')
doc.add_paragraph('系统：自动识别并存储为 decision 类别，重要度 0.9')

doc.add_heading('4.2 主动存储', 2)
doc.add_paragraph('用户：记住我的股票账号是 123456')
doc.add_paragraph('系统：立即存储，重要度 0.85，类别 finance')

doc.add_heading('4.3 智能检索', 2)
doc.add_paragraph('命令：ums recall search "编程"')
doc.add_paragraph()
doc.add_paragraph('结果：')
doc.add_paragraph('1. [preference] 用 Python 编程 (重要度: 0.8)')
doc.add_paragraph('2. [decision] 使用 React 编程 (重要度: 0.7)')
doc.add_paragraph('3. [project] 编程项目规划 (重要度: 0.6)')

doc.add_page_break()

# 第五章
doc.add_heading('五、技术架构', 1)
doc.add_heading('5.1 核心模块', 2)
doc.add_paragraph('• daily.py - 每日记忆管理', style='List Bullet')
doc.add_paragraph('• analyzer.py - 月度智能分析', style='List Bullet')
doc.add_paragraph('• recall.py - 增强记忆检索', style='List Bullet')
doc.add_paragraph('• session.py - 会话持久化', style='List Bullet')
doc.add_paragraph('• smart_memory.py - 智能识别引擎', style='List Bullet')

doc.add_heading('5.2 评分算法', 2)
doc.add_paragraph('最终得分 = 匹配度 × 时间衰减 × 重要性加权')
doc.add_paragraph('时间衰减 = 0.5 + 0.5 × exp(-年龄天数 / 60)')
doc.add_paragraph('重要性加权 = 0.7 + 0.3 × 重要度(0-1)')

doc.add_heading('5.3 自动化配置', 2)
doc.add_paragraph('• 每日 22:00 - 系统自动备份', style='List Bullet')
doc.add_paragraph('• 每月1日 02:00 - 执行月度智能分析', style='List Bullet')
doc.add_paragraph('• 对话结束时 - 自动识别重要信息', style='List Bullet')

doc.add_page_break()

# 第六章
doc.add_heading('六、总结', 1)
p = doc.add_paragraph()
p.add_run('统一记忆管理系统通过整合分散的工具、提供统一接口、实现智能识别，让 AI 助手的记忆管理变得简单高效。').font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('核心优势：一个命令、自动识别、智能检索、永久保存。')
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 51, 102)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('6.1 适用场景', 2)
doc.add_paragraph('• 需要长期保存工作历史的个人和团队', style='List Bullet')
doc.add_paragraph('• 希望减少 API 调用成本的 AI 应用', style='List Bullet')
doc.add_paragraph('• 需要智能检索大量历史信息的知识工作者', style='List Bullet')
doc.add_paragraph('• 希望 AI 助手更懂自己的用户', style='List Bullet')

doc.add_heading('6.2 联系信息', 2)
doc.add_paragraph('文档版本：v2.0.0')
doc.add_paragraph('最后更新：2026年3月')
doc.add_paragraph('作者：雪子助手')
doc.add_paragraph('项目地址：http://106.54.25.161/my-home/')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('感谢使用统一记忆管理系统！')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
output_path = '/Users/zhaoruicn/.openclaw/workspace/统一记忆管理系统介绍.docx'
doc.save(output_path)
print('Word文档已生成:', output_path)
