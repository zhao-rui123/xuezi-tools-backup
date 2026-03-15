#!/usr/bin/env python3
"""
文档+图表一体化引擎
支持在Word文档中直接嵌入图表
"""

import os
import io
import json
import pandas as pd
from typing import Dict, List, Optional, Union, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')

from core.smart_chart_recommender import SmartChartRecommender, recommend_chart, ChartType


class DocumentWithCharts:
    def __init__(self, template_path: Optional[str] = None):
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        self.charts: List[Dict] = []
        self.data_sources: Dict[str, pd.DataFrame] = {}
        self._chart_counter = 0
        self._set_default_styles()

    def _set_default_styles(self):
        self.styles = {
            "title": {"font_size": 22, "bold": True, "color": RGBColor(44, 62, 80)},
            "heading1": {"font_size": 16, "bold": True, "color": RGBColor(44, 62, 80)},
            "heading2": {"font_size": 14, "bold": True, "color": RGBColor(52, 73, 94)},
            "heading3": {"font_size": 12, "bold": True, "color": RGBColor(67, 73, 84)},
            "normal": {"font_size": 11, "color": RGBColor(44, 62, 80)},
            "caption": {"font_size": 9, "color": RGBColor(127, 140, 141)}
        }

    def add_title(self, text: str, level: int = 0) -> None:
        heading = self.doc.add_heading(text, level)
        for run in heading.runs:
            if level == 0:
                run.font.size = Pt(self.styles["title"]["font_size"])
                run.font.bold = True
                run.font.color.rgb = self.styles["title"]["color"]
            elif level == 1:
                run.font.size = Pt(self.styles["heading1"]["font_size"])
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(self.styles["heading2"]["font_size"])
                run.font.bold = True

    def add_paragraph(self, text: str, style: str = "normal") -> None:
        para = self.doc.add_paragraph(text)
        for run in para.runs:
            run.font.size = Pt(self.styles[style]["font_size"])
            run.font.color.rgb = self.styles[style]["color"]
        return para

    def add_chart(self, chart_type: str, data: Union[Dict, pd.DataFrame, str],
                  title: str = None, width: float = 6, height: float = 4,
                  position: str = "auto", auto_recommend: bool = True) -> str:
        self._chart_counter += 1
        chart_id = f"chart_{self._chart_counter}"

        if isinstance(data, str):
            if data.endswith('.csv'):
                df = pd.read_csv(data)
            elif data.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(data)
            elif data.endswith('.json'):
                df = pd.read_json(data)
            else:
                raise ValueError(f"Unsupported file format: {data}")
        elif isinstance(data, dict):
            df = pd.DataFrame(data)
        elif isinstance(data, pd.DataFrame):
            df = data
        else:
            raise ValueError("Unsupported data format")

        self.data_sources[chart_id] = df

        if auto_recommend and chart_type == "auto":
            recommendation = recommend_chart(df)
            chart_type = recommendation["recommended_chart"]
            if not title:
                title = recommendation["config"].get("title", "图表")

        chart_path = self._generate_chart_image(chart_type, df, chart_id, title)

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_picture(chart_path, width=Inches(width), height=Inches(height))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if title:
            caption = self.doc.add_paragraph(title, style="caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.charts.append({
            "id": chart_id,
            "type": chart_type,
            "title": title,
            "path": chart_path,
            "data": df.to_dict() if isinstance(df, pd.DataFrame) else df
        })

        return chart_id

    def _generate_chart_image(self, chart_type: str, df: pd.DataFrame,
                            chart_id: str, title: str = None) -> str:
        fig, ax = plt.subplots(figsize=(8, 5))

        columns = df.columns.tolist()
        if len(columns) < 2:
            ax.bar(range(len(df)), df[columns[0]]) if len(columns) == 1 else ax.bar([0], [0])
            ax.set_xticks(range(len(df)))
            ax.set_xticklabels(df[columns[0]].tolist() if len(columns) == 1 else ["数据"])
        else:
            x_col = columns[0]
            y_cols = columns[1:]

            x_data = df[x_col].tolist()
            x_pos = range(len(x_data))

            colors = ['#4A90E2', '#E24A4A', '#2ECC71', '#F39C12', '#9B59B6']

            if chart_type in ["bar", "BAR"]:
                for i, y_col in enumerate(y_cols):
                    ax.bar([x + i*0.8/len(y_cols) for x in x_pos],
                          df[y_col], width=0.8/len(y_cols),
                          label=y_col, color=colors[i % len(colors)])
                ax.set_xticks(x_pos)
                ax.set_xticklabels(x_data, rotation=45)

            elif chart_type in ["line", "LINE"]:
                for i, y_col in enumerate(y_cols):
                    ax.plot(x_data, df[y_col], marker='o',
                           label=y_col, color=colors[i % len(colors)], linewidth=2)
                ax.set_xticks(x_pos)
                ax.set_xticklabels(x_data, rotation=45)

            elif chart_type in ["pie", "PIE"]:
                ax.pie(df[y_cols[0]], labels=df[x_col], autopct='%1.1f%%',
                      colors=colors[:len(df)], startangle=90)

            elif chart_type in ["horizontal_bar", "HORIZONTAL_BAR"]:
                y_pos = range(len(x_data))
                for i, y_col in enumerate(y_cols):
                    ax.barh(y_pos, df[y_col], height=0.6,
                           label=y_col, color=colors[i % len(colors)])
                ax.set_yticks(y_pos)
                ax.set_yticklabels(x_data)

            elif chart_type in ["area", "AREA"]:
                for i, y_col in enumerate(y_cols):
                    ax.fill_between(x_data, df[y_col], alpha=0.5,
                                   label=y_col, color=colors[i % len(colors)])
                ax.set_xticks(x_pos)
                ax.set_xticklabels(x_data, rotation=45)

            else:
                for i, y_col in enumerate(y_cols):
                    ax.bar([x + i*0.8/len(y_cols) for x in x_pos],
                          df[y_col], width=0.8/len(y_cols),
                          label=y_col, color=colors[i % len(colors)])
                ax.set_xticks(x_pos)
                ax.set_xticklabels(x_data, rotation=45)

        if title:
            ax.set_title(title, fontsize=14, fontweight='bold', pad=10)
        ax.legend(loc='best', fontsize=9)
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        plt.tight_layout()

        temp_dir = os.path.join(os.path.dirname(__file__), '..', 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        chart_path = os.path.join(temp_dir, f"{chart_id}.png")

        plt.savefig(chart_path, dpi=150, bbox_inches='tight',
                   facecolor='white', edgecolor='none')
        plt.close(fig)

        return chart_path

    def add_table(self, data: Union[Dict, List, pd.DataFrame],
                  headers: List[str] = None, style: str = "Light Grid Accent 1") -> None:
        if isinstance(data, pd.DataFrame):
            df = data
            if headers is None:
                headers = df.columns.tolist()
            table_data = [headers] + df.values.tolist()
        elif isinstance(data, dict):
            if headers is None:
                headers = list(data.keys())
            table_data = [headers] + [[row.get(h, "") for h in headers] for row in data.values()]
        else:
            table_data = data if headers else [[]]

        if not table_data:
            return

        table = self.doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = style

        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

    def add_image(self, image_path: str, width: float = 6, caption: str = None) -> None:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image not found: {image_path}")

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if caption:
            caption_para = self.doc.add_paragraph(caption, style="caption")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    def add_section(self, title: str, content: str = None) -> None:
        self.add_title(title, level=1)
        if content:
            self.add_paragraph(content)

    def save(self, output_path: str) -> None:
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        self.doc.save(output_path)
        print(f"文档已生成: {output_path}")

    def get_summary(self) -> Dict:
        return {
            "charts_count": len(self.charts),
            "charts": [c["type"] for c in self.charts],
            "data_sources": list(self.data_sources.keys())
        }


def create_data_driven_report(data_source: str, output_path: str,
                             chart_configs: List[Dict] = None,
                             report_title: str = "数据报告") -> str:
    if data_source.endswith('.csv'):
        df = pd.read_csv(data_source)
    elif data_source.endswith(('.xlsx', '.xls')):
        df = pd.read_excel(data_source)
    elif data_source.endswith('.json'):
        df = pd.read_json(data_source)
    else:
        raise ValueError(f"Unsupported file format: {data_source}")

    doc = DocumentWithCharts()
    doc.add_title(report_title, level=0)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_page_break()

    doc.add_title("数据概览", level=1)
    doc.add_paragraph(f"数据行数: {len(df)}, 数据列数: {len(df.columns)}")

    doc.add_title("数据预览", level=2)
    doc.add_table(df.head(10))

    if chart_configs:
        for config in chart_configs:
            chart_type = config.get("type", "auto")
            title = config.get("title")
            columns = config.get("columns")

            if columns:
                chart_data = df[columns]
            else:
                chart_data = df

            doc.add_chart(chart_type, chart_data, title=title)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    sample_data = {
        "month": ["1月", "2月", "3月", "4月", "5月", "6月"],
        "销售额": [12000, 15000, 18000, 16000, 20000, 25000],
        "利润": [3000, 4500, 5500, 4800, 6500, 8000]
    }

    doc = DocumentWithCharts()
    doc.add_title("月度销售报告", level=0)
    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_section("销售数据", "以下是本月销售数据摘要：")
    doc.add_table(sample_data)

    doc.add_section("销售趋势", "销售额和利润趋势图：")
    doc.add_chart("line", sample_data, title="月度销售趋势", width=6, height=4)

    doc.add_section("销售对比", "各月销售额对比：")
    doc.add_chart("bar", sample_data, title="月度销售额", width=6, height=4)

    doc.save("sample_report.docx")
    print("报告生成完成: sample_report.docx")
