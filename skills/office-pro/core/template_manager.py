#!/usr/bin/env python3
"""
模板管理系统
支持报告模板、图表模板和组合模板
"""

import os
import json
import shutil
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from docx import Document
from openpyxl import Workbook
from pptx import Presentation
import pandas as pd


@dataclass
class Template:
    name: str
    category: str
    description: str
    file_path: str
    template_type: str
    config: Dict = field(default_factory=dict)
    preview: str = None


class TemplateManager:
    TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')

    REPORT_CATEGORIES = {
        "financial": "财务报告",
        "marketing": "营销报告",
        "project": "项目报告",
        "sales": "销售报告",
        "hr": "人力资源",
        "academic": "学术论文"
    }

    CHART_CATEGORIES = {
        "business": "商业图表",
        "scientific": "科学图表",
        "infographic": "信息图表",
        "dashboard": "仪表盘"
    }

    def __init__(self, template_dir: str = None):
        self.template_dir = template_dir or self.TEMPLATE_DIR
        self.templates: List[Template] = []
        self._load_templates()

    def _load_templates(self):
        if not os.path.exists(self.template_dir):
            return

        for root, dirs, files in os.walk(self.template_dir):
            for file in files:
                if file.endswith(('.docx', '.xlsx', '.pptx', '.json')):
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, self.template_dir)
                    category = rel_path.split(os.sep)[0] if os.sep in rel_path else "other"

                    config = {}
                    if file.endswith('.json'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            config = json.load(f)
                        template_type = config.get('type', 'config')
                    elif file.endswith('.docx'):
                        template_type = 'word'
                    elif file.endswith('.xlsx'):
                        template_type = 'excel'
                    elif file.endswith('.pptx'):
                        template_type = 'powerpoint'

                    template = Template(
                        name=config.get('name', file.replace('.json', '')),
                        category=category,
                        description=config.get('description', ''),
                        file_path=file_path,
                        template_type=template_type,
                        config=config,
                        preview=config.get('preview')
                    )
                    self.templates.append(template)

    def list_templates(self, category: str = None, template_type: str = None) -> List[Template]:
        result = self.templates

        if category:
            result = [t for t in result if t.category == category]

        if template_type:
            result = [t for t in result if t.template_type == template_type]

        return result

    def get_template(self, name: str) -> Optional[Template]:
        for template in self.templates:
            if template.name == name:
                return template
        return None

    def apply_report_template(self, template_name: str, data: Dict,
                              output_path: str) -> str:
        template = self.get_template(template_name)
        if not template:
            raise ValueError(f"Template not found: {template_name}")

        if template.template_type == 'word':
            return self._apply_word_template(template, data, output_path)
        elif template.template_type == 'excel':
            return self._apply_excel_template(template, data, output_path)
        else:
            raise ValueError(f"Unsupported template type: {template.template_type}")

    def _apply_word_template(self, template: Template, data: Dict, output_path: str) -> str:
        if os.path.exists(template.file_path):
            doc = Document(template.file_path)
        else:
            doc = Document()

        for para in doc.paragraphs:
            for key, value in data.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        doc.save(output_path)
        return output_path

    def _apply_excel_template(self, template: Template, data: Dict, output_path: str) -> str:
        if os.path.exists(template.file_path):
            wb = load_workbook(template.file_path)
        else:
            wb = Workbook()

        ws = wb.active

        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for key, value in data.items():
                        if key in cell.value:
                            cell.value = cell.value.replace(key, str(value))

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        wb.save(output_path)
        return output_path

    def create_combination_template(self, name: str, category: str,
                                   description: str, components: List[Dict],
                                   output_path: str) -> str:
        config = {
            "name": name,
            "category": category,
            "description": description,
            "type": "combination",
            "components": components
        }

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

        return output_path


class TemplateBuilder:
    def __init__(self):
        self.components: List[Dict] = []
        self.name = ""
        self.category = ""
        self.description = ""

    def set_name(self, name: str) -> "TemplateBuilder":
        self.name = name
        return self

    def set_category(self, category: str) -> "TemplateBuilder":
        self.category = category
        return self

    def set_description(self, description: str) -> "TemplateBuilder":
        self.description = description
        return self

    def add_section(self, title: str, content_type: str = "text",
                   data_source: str = None, chart_config: Dict = None) -> "TemplateBuilder":
        component = {
            "type": "section",
            "title": title,
            "content_type": content_type,
            "data_source": data_source,
            "chart_config": chart_config or {}
        }
        self.components.append(component)
        return self

    def add_chart_placeholder(self, title: str, chart_type: str = "auto",
                            data_source: str = None) -> "TemplateBuilder":
        component = {
            "type": "chart",
            "title": title,
            "chart_type": chart_type,
            "data_source": data_source
        }
        self.components.append(component)
        return self

    def add_table_placeholder(self, title: str, data_source: str = None,
                             rows: int = 10) -> "TemplateBuilder":
        component = {
            "type": "table",
            "title": title,
            "data_source": data_source,
            "rows": rows
        }
        self.components.append(component)
        return self

    def build(self, output_path: str) -> str:
        config = {
            "name": self.name,
            "category": self.category,
            "description": self.description,
            "type": "combination",
            "components": self.components
        }

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

        return output_path


def register_template(template_path: str, name: str, category: str,
                     description: str, template_type: str) -> Template:
    config = {
        "name": name,
        "category": category,
        "description": description,
        "type": template_type
    }

    config_path = template_path.replace('.docx', '.json').replace('.xlsx', '.json').replace('.pptx', '.json')

    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

    return Template(
        name=name,
        category=category,
        description=description,
        file_path=template_path,
        template_type=template_type,
        config=config
    )


if __name__ == "__main__":
    manager = TemplateManager()

    print("可用模板:")
    print("=" * 50)

    print("\n报告模板:")
    for template in manager.list_templates(template_type='word'):
        print(f"  - {template.name} ({template.category}): {template.description}")

    print("\n组合模板:")
    for template in manager.list_templates():
        if template.template_type == 'config':
            print(f"  - {template.name}: {template.description}")
