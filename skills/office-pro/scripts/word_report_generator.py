#!/usr/bin/env python3
"""
Word 报告生成器
支持模板驱动和数据填充
"""

import argparse
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os


def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold


def create_report(data, output_path, template_path=None):
    """创建 Word 报告"""
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    
    # 标题
    title = doc.add_heading(data.get('title', '报告'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}")
    set_chinese_font(date_run, font_size=10)
    
    doc.add_paragraph()  # 空行
    
    # 章节
    for section in data.get('sections', []):
        # 章节标题
        heading = doc.add_heading(section.get('title', ''), level=1)
        
        # 内容
        for content in section.get('content', []):
            if isinstance(content, str):
                para = doc.add_paragraph(content)
            elif isinstance(content, dict):
                if content.get('type') == 'list':
                    for item in content.get('items', []):
                        doc.add_paragraph(item, style='List Bullet')
                elif content.get('type') == 'table':
                    table_data = content.get('data', [])
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                table.rows[i].cells[j].text = str(cell_data)
    
    # 保存
    doc.save(output_path)
    print(f"报告已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Word 报告生成器')
    parser.add_argument('--data', '-d', required=True, help='数据 JSON 文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--template', '-t', help='模板文件路径')
    
    args = parser.parse_args()
    
    # 读取数据
    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    create_report(data, args.output, args.template)


if __name__ == '__main__':
    main()
