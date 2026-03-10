#!/usr/bin/env python3
"""
PCS竞品分析报告生成器
基于IES920和BCS系列技术规格对比
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = docx.oxml.OxmlElement(tag)
                    tcPr.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), edge_data.get('color', '000000'))

def create_pcs_comparison_report():
    """创建PCS竞品分析报告"""
    
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 标题
    title = doc.add_heading('储能变流器（PCS）竞品分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'IES920系列 vs BCS系列技术规格对比\n报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)
    
    doc.add_paragraph()
    
    # 1. 报告概述
    doc.add_heading('一、报告概述', 1)
    doc.add_paragraph(
        '本报告针对两款主流储能变流器（PCS）产品进行技术规格对比分析，'
        '为PCS产品研发提供参考依据。对比产品包括：'
    )
    
    # 产品列表
    products = [
        ('IES920系列', 'IES920-07-3500-C-D / IES920-07-3150-C-D'),
        ('BCS系列', 'BCS3500K-C-HUD / BCS3150K-C-HUD')
    ]
    
    for i, (name, model) in enumerate(products, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{name}：').bold = True
        p.add_run(model)
    
    doc.add_paragraph()
    
    # 2. 技术规格对比表
    doc.add_heading('二、技术规格详细对比', 1)
    
    # 2.1 直流侧参数对比
    doc.add_heading('2.1 直流侧参数', 2)
    
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['参数项目', 'IES920-3500', 'IES920-3150', 'BCS系列']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # 数据行
    dc_data = [
        ['最大直流电压', '1500V', '1500V', '1500Vdc'],
        ['直流工作电压范围', '1000~1500V', '1000~1500V', '1000~1500 Vdc'],
        ['最大直流电流', '1964×2 A', '1767×2 A', '3536A / 2*1964A'],
        ['直流接入支数', '2路', '2路', '1/2路']
    ]
    
    for i, row_data in enumerate(dc_data, 1):
        for j, text in enumerate(row_data):
            table1.rows[i].cells[j].text = text
    
    doc.add_paragraph()
    
    # 2.2 交流侧参数对比（并网）
    doc.add_heading('2.2 交流侧参数（并网模式）', 2)
    
    table2 = doc.add_table(rows=8, cols=4)
    table2.style = 'Light Grid Accent 1'
    
    headers2 = ['参数项目', 'IES920-3500', 'IES920-3150', 'BCS系列']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    ac_data = [
        ['额定输出功率', '3500kW@45°C', '3150kW@50°C', '3500kW / 3150kW'],
        ['最大输出功率', '3850kW@35°C', '3465kW@40°C', '3850kVA / 3465kVA'],
        ['额定电网电压', '690V AC', '690V AC', '690Vac'],
        ['允许电网电压范围', '621~759V', '621~759V', '-10%~+10%（可设置）'],
        ['电网频率', '50Hz', '50Hz', '50Hz/60Hz（可设置）'],
        ['最大输出电流', '3222A@35°C', '2899A@40°C', '3222A / 2900A'],
        ['THDi（电流谐波）', '<3%', '<3%', '<3%']
    ]
    
    for i, row_data in enumerate(ac_data, 1):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text
    
    doc.add_paragraph()
    
    # 2.3 效率与性能
    doc.add_heading('2.3 效率与性能指标', 2)
    
    table3 = doc.add_table(rows=7, cols=4)
    table3.style = 'Light Grid Accent 1'
    
    headers3 = ['参数项目', 'IES920系列', 'BCS系列', '对比分析']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    perf_data = [
        ['最大效率', '未明确标注', '>99%', 'BCS明确标注>99%'],
        ['功率因数', '-105%~+105%', '-105%~+105%', '相同'],
        ['充放电转换时间', '<30ms', '未明确', 'IES920指标更优'],
        ['电压精度（离网）', '<1%U_N', '1%', '相当'],
        ['电压失真度', '<3%', '<2%', 'BCS略优'],
        ['过载能力', '未明确', '110%', 'BCS明确过载能力']
    ]
    
    for i, row_data in enumerate(perf_data, 1):
        for j, text in enumerate(row_data):
            table3.rows[i].cells[j].text = text
    
    doc.add_paragraph()
    
    # 2.4 物理参数
    doc.add_heading('2.4 物理参数与环境适应性', 2)
    
    table4 = doc.add_table(rows=9, cols=3)
    table4.style = 'Light Grid Accent 1'
    
    headers4 = ['参数项目', 'IES920系列', 'BCS系列']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    physical_data = [
        ['尺寸（宽×高×深）', '未标注', '1200×2450×1250mm'],
        ['重量', '未标注', '1700kg'],
        ['防护等级', '未明确', 'IP65'],
        ['防腐等级', '未明确', 'C4（C5选配）'],
        ['冷却方式', '未明确', '智能液冷'],
        ['最高工作海拔', '未明确', '5000m（>3000m降额）'],
        ['工作环境温度', '未明确', '-35℃~60℃（>50℃降额）'],
        ['噪声', '<85dB', '未标注']
    ]
    
    for i, row_data in enumerate(physical_data, 1):
        for j, text in enumerate(row_data):
            table4.rows[i].cells[j].text = text
    
    doc.add_paragraph()
    
    # 2.5 通信与标准
    doc.add_heading('2.5 通信协议与标准', 2)
    
    table5 = doc.add_table(rows=5, cols=3)
    table5.style = 'Light Grid Accent 1'
    
    headers5 = ['项目', 'IES920系列', 'BCS系列']
    for i, header in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    comm_data = [
        ['通信协议', '未明确标注', 'Modbus-RTU / Modbus-TCP / IEC61850 / IEC104'],
        ['通信方式', '未明确标注', 'RS485/CAN/Ethernet'],
        ['隔离方式', '无隔离', '无隔离变压器'],
        ['符合标准', '未明确标注', 'GB/T 34120、GB/T 34133']
    ]
    
    for i, row_data in enumerate(comm_data, 1):
        for j, text in enumerate(row_data):
            table5.rows[i].cells[j].text = text
    
    doc.add_page_break()
    
    # 3. 核心差异分析
    doc.add_heading('三、核心差异分析', 1)
    
    doc.add_heading('3.1 IES920系列优势', 2)
    advantages_ies = [
        '充放电转换时间<30ms，响应速度更快',
        '噪声控制<85dB，声环境友好',
        '详细的电流参数标注（1964×2 / 1767×2）'
    ]
    for adv in advantages_ies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(adv)
    
    doc.add_heading('3.2 BCS系列优势', 2)
    advantages_bcs = [
        '效率指标明确：最大效率>99%',
        '物理参数完整：尺寸、重量、防护等级清晰',
        '环境适应性强：-35℃~60℃工作温度，5000m海拔',
        '冷却方式先进：智能液冷',
        '通信协议丰富：支持IEC61850等工业标准',
        '过载能力明确：110%过载能力',
        '防护等级高：IP65',
        '符合国家标准：GB/T 34120、GB/T 34133'
    ]
    for adv in advantages_bcs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(adv)
    
    doc.add_heading('3.3 信息完整度对比', 2)
    doc.add_paragraph(
        'BCS系列产品规格书信息完整度明显高于IES920系列，'
        '包括物理尺寸、重量、环境适应性、冷却方式、通信协议等关键参数均有详细标注。'
        'IES920系列在部分性能指标上存在信息缺失。'
    )
    
    doc.add_page_break()
    
    # 4. 研发建议
    doc.add_heading('四、PCS研发建议', 1)
    
    doc.add_heading('4.1 技术参数目标', 2)
    targets = [
        ('直流侧', '1500V电压等级，1000-1500V工作范围，双支路接入'),
        ('交流侧', '690V输出，50Hz/60Hz兼容，<3%谐波'),
        ('功率等级', '3500kW/3150kW双规格'),
        ('效率', '目标>99%，明确标注最大效率'),
        ('响应速度', '充放电转换<30ms'),
        ('过载能力', '明确110%过载能力')
    ]
    
    for category, target in targets:
        p = doc.add_paragraph()
        p.add_run(f'{category}：').bold = True
        p.add_run(target)
    
    doc.add_heading('4.2 环境适应性设计', 2)
    env_design = [
        '工作温度范围：-35℃~60℃，>50℃自动降额',
        '防护等级：IP65（防尘防水）',
        '防腐等级：C4（沿海）/ C5（C5选配，高腐蚀环境）',
        '海拔适应：5000m，>3000m自动降额',
        '冷却方式：智能液冷（优于风冷）'
    ]
    for item in env_design:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_heading('4.3 通信与标准', 2)
    comm_design = [
        '支持Modbus-RTU/TCP、IEC61850、IEC104',
        '通信接口：RS485、CAN、Ethernet',
        '符合GB/T 34120、GB/T 34133国家标准',
        '预留远程监控接口'
    ]
    for item in comm_design:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_heading('4.4 产品规格书完善', 2)
    doc.add_paragraph(
        '建议参考BCS系列的完整性，产品规格书应包含以下关键信息：'
    )
    spec_items = [
        '详细的物理尺寸和重量',
        '完整的环境适应性参数',
        '明确的效率指标',
        '过载能力说明',
        '通信协议详细列表',
        '符合的标准清单',
        '噪声指标'
    ]
    for item in spec_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_page_break()
    
    # 5. 结论
    doc.add_heading('五、结论', 1)
    
    doc.add_paragraph(
        '通过对IES920系列和BCS系列储能变流器的技术规格对比分析，'
        '两款产品在核心电气参数（1500Vdc、690Vac、3500kW/3150kW）方面基本一致，'
        '均可作为PCS研发的参考标杆。'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'BCS系列在产品规格完整性、环境适应性、通信协议丰富度等方面表现更优，'
        '建议以其为主要参考对象进行PCS产品设计。'
        '同时，IES920系列在充放电转换速度和噪声控制方面的优势也值得借鉴。'
    )
    
    doc.add_paragraph()
    
    conclusion_points = [
        '技术路线：1500V高压直流+690V交流，3500kW大功率',
        '冷却方案：推荐智能液冷',
        '防护等级：IP65',
        '效率目标：>99%',
        '标准符合：GB/T 34120、GB/T 34133'
    ]
    
    for point in conclusion_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point)
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 报告结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    output_path = '/Users/zhaoruicn/.openclaw/workspace/PCS竞品分析报告.docx'
    doc.save(output_path)
    print(f'✅ 报告已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_pcs_comparison_report()
