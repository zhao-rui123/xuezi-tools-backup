#!/usr/bin/env python3
"""生成主控制台体系介绍文档（从修复飞书卡片开始·脱敏版）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
title = doc.add_heading('从零搭建飞书主控制台体系', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'创建时间：{datetime.date.today().strftime("%Y-%m-%d")}')
doc.add_paragraph('本文档记录了如何从「飞书卡片按钮点不通」这个问题出发，'
                  '一步步搭建起一套完整的 AI 助手主控制台体系。所有敏感信息均已脱敏。')

# 背景
doc.add_heading('一、从一个问题开始：飞书卡片按钮点不通', level=2)
doc.add_paragraph('AI 助手通过 OpenClaw 向飞书发送交互卡片，'
                  '卡片里有按钮（action buttons），用户点按钮后应该能触发相应命令。')
doc.add_paragraph('但实际使用时发现：按钮点击后没有任何反应。')
doc.add_paragraph('这看起来是一个「按钮失效」问题，但真正的原因和解决方案比想象的要深。')

# 根因分析
doc.add_heading('二、根因：按钮 value 格式错误', level=2)
doc.add_paragraph('OpenClaw Feishu 插件已开启 inlineButtons=dm，说明按钮事件确实能到达 gateway。')
doc.add_paragraph('排查后发现：按钮的 value 发成了 JSON 字符串，而不是结构化对象。')

p = doc.add_paragraph()
p.add_run('错误做法（JSON 字符串）：').bold = True
doc.add_paragraph('{"oc":"ocf1","k":"quick","q":"/help","c":{...}}')
p = doc.add_paragraph()
p.add_run('正确做法（OCF1 Structured Object）：').bold = True
code = '''{
  "oc": "ocf1",
  "k": "quick",
  "a": "feishu.quick_actions.xxx",
  "q": "实际命令文字",
  "c": {
    "u": "用户 Open ID",
    "e": 1893456000000,
    "t": "p2p"
  }
}'''
doc.add_paragraph(code)

# 飞书卡片协议
doc.add_heading('三、飞书 OCF1 Structured Quick Action 协议', level=2)
doc.add_paragraph('飞书支持一种叫 OCF1 的结构化按钮协议，格式如下：')

fields = [
    ('oc', '"ocf1"', '固定值，表示这是 OCF1 结构化按钮'),
    ('k', '"quick"', '固定值，表示这是 quick action'),
    ('a', 'feishu.quick_actions.xxx', 'action 路由名，决定按钮事件发给哪个 handler'),
    ('q', '实际命令文字', '用户可见的命令文字，用于路由匹配'),
    ('c', '{u, e, t}', '用户上下文：u=OpenID，e=过期时间，t=聊天类型'),
]
for name, example, desc in fields:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(f'{example} — {desc}')

doc.add_paragraph('当用户点按钮，gateway 收到命令后，根据 q 和 a 路由到对应的处理器函数。')

# 系统演进
doc.add_heading('四、系统演进：从展示到可操作', level=2)

doc.add_heading('阶段 1：把展示系统做出来', level=3)
doc.add_paragraph('首先把各种状态的专题面板做成飞书卡片，包括：')
panels = [
    '系统面板（健康检查/备份/Gateway/Tailscale/云服务器）',
    '任务面板（阻塞/高优/收件箱/自动导入任务）',
    '模型面板（OpenClaw 会话模型 + Claude Code 模型）',
    '记忆面板（候选人统计/风险/规则/TODO）',
    '定时任务中心（每日备份/云同步/股票推送/会话快照）',
    '后台任务中心（当前模型/上下文/活跃任务）',
    '执行中心（ACP/screen/后台任务统一视图）',
]
for name in panels:
    doc.add_paragraph(f'• {name}')

doc.add_heading('阶段 2：从数字到语义', level=3)
doc.add_paragraph('用户看到的应该是「任务本身」，不是技术 id。')
doc.add_paragraph('于是把任务名翻成业务名：')
examples = [
    '"cc-min-test" → "Claude Code 最小链路测试"',
    '"cc-wrapper-test" + workdir 含 railway-storage → "铁路储能 5MWh 项目（CC包装链路测试）"',
]
for ex in examples:
    doc.add_paragraph(f'• {ex}')
doc.add_paragraph('同时给每个任务附上链路 / 状态 / 最近时间，统一格式：「任务名｜执行方式｜状态｜最近时间」。')

doc.add_heading('阶段 3：主控制台 V3 — 语义化首页', level=3)
doc.add_paragraph('首页不只是导航，开始显示「现在在干什么、哪里有问题、该点哪里」：')
doc.add_paragraph('• 顶部 6 个关键指标字段（模型/阻塞数/Inbox/风险/备份状态）')
doc.add_paragraph('• 「当前最重要的事」— Top1 优先 + 所有高优先级建议')
doc.add_paragraph('• 「当前活跃/最近任务」— 业务任务名 + 链路 + 状态')

doc.add_heading('阶段 4：主控制台 V4 — 控制入口更显眼', level=3)
doc.add_paragraph('在首页 Action 区加入「轻控制入口」按钮，更显眼。')
doc.add_paragraph('同时把 Top1 优先从列表改为单行显示，一眼就知道最该干什么。')

doc.add_heading('阶段 5：主控制台 V5 — 两级导航', level=3)
doc.add_paragraph('发现首页内容太多，飞书卡片有限制（约 4KB/条），完整首页超 11KB。')
doc.add_paragraph('解决方案：改成两级导航。')
doc.add_paragraph('• Hub 首页（V5）：5个一级入口，每个入口下面有二级标题说明（文字，无需点击）')
doc.add_paragraph('• 专题面板（二级）：点进去后是该专题下的所有具体按钮')

hubs = [
    ('⚡ 执行链路', '执行中心 | 语义任务面板 | ACP | screen详情 | 重跑测试任务 | 任务历史'),
    ('📋 任务与记忆', '任务面板 | 记忆面板 | 今日重点 | 阻塞 | 高优先级 | Inbox'),
    ('🩺 系统与定时', '系统面板 | 定时任务中心 | 健康检查 | 备份状态'),
    ('🤖 模型与控制', '模型面板 | 切换CC模型 | 会话状态 | 轻控制入口 | 清理测试任务'),
    ('🔧 快捷', '快捷动作面板 | 主控制台 | 帮助'),
]
for name, subs in hubs:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(subs)

# 控制动作
doc.add_heading('五、真实可执行的控制闭环示例：重跑测试任务', level=2)
doc.add_paragraph('到这一步，系统已经从「展示系统」升级为「可操作系统」。')

steps = [
    ('第1步：候选面板', '列出 3 个白名单测试任务（cc-min-test / cc-wrapper-test / codex-wrapper-test），每个任务一个「确认重跑」按钮。'),
    ('第2步：二次确认卡', '点完按钮后弹出确认卡，显示任务详情（工作目录/日志/上次时间）。必须再次点「确认并重跑」才触发真执行。'),
    ('第3步：命令处理器', '按钮发送「rerun execute {task}」，handle_rerun_cmd.py 校验白名单后调用 rerun-test-task.sh。'),
    ('第4步：真实执行', 'rerun-test-task.sh 复用原始 AGENT/WORKDIR/PROMPT_FILE，生成新任务名（xxx-rerun-MMDD-HHMMSS）。'),
    ('第5步：结果卡', 'rerun-test-task-result.sh 输出标准化结果，rerun_result_card.py 包装成结果卡发回用户。'),
]
for name, desc in steps:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(desc)

# 控制动作分级
doc.add_heading('六、控制动作风险分级', level=2)
levels = [
    ('P0（低风险）', '现在就能开', '查看详情/跳转/返回/查看状态/历史/最近时间'),
    ('P1（中风险）', '可做但需确认', '清理测试任务/重跑测试任务/停止明确标记的测试任务'),
    ('P2（高风险）', '后面再做', '停止正式项目任务/重跑正式项目任务/恢复失败任务/改定时任务配置'),
]
for name, level, items in levels:
    p = doc.add_paragraph()
    p.add_run(f'{name}（{level}）：').bold = True
    p.add_run(items)

# 关键文件
doc.add_heading('七、关键文件清单', level=2)
files = [
    ('cockpit_v5_hub.py', 'Hub 首页生成脚本（带二级标题文字）'),
    ('cockpit_v5.py', 'Hub + 各专题二级面板生成脚本'),
    ('cockpit-v5-hub.json', 'Hub 首页卡片 JSON（精简版，11KB → 4KB 内）'),
    ('cockpit-v5-{key}.json', '各专题二级面板卡片 JSON'),
    ('handle_rerun_cmd.py', '重跑命令处理器（白名单校验）'),
    ('rerun-test-task.sh', '真实重跑脚本（只允许白名单任务）'),
    ('rerun-test-task-result.sh', '重跑结果回显脚本'),
    ('rerun_result_card.py', '重跑结果卡生成器'),
    ('rerun_test_tasks_panel_card.py', '重跑候选面板生成器'),
    ('rerun_test_confirm_card.py', '重跑二次确认卡生成器'),
    ('model_panel_card.py', '模型面板（含 OpenClaw 4 个模型 + CC 3 个模型切换）'),
    ('control_actions_policy_card.py', '控制动作分级卡片'),
    ('light_control_panel_card.py', '轻控制入口面板'),
    ('docs/cockpit-v5.md', '主控制台 V5 导航文档'),
    ('docs/cockpit-index.md', '完整面板/脚本/路由索引'),
]
for fname, desc in files:
    p = doc.add_paragraph()
    p.add_run(f'• {fname}').bold = True
    p.add_run(f' — {desc}')

# 设计原则
doc.add_heading('八、核心设计原则', level=2)
principles = [
    ('结果必须可反馈', '执行后必须回显结果，失败也要有明确提示，不能静默。'),
    ('低风险优先', '先做 P0，P1 试点，P2 暂缓。不一上来就开放高危动作。'),
    ('无重复设计', '同一专题只出现在一个入口下，用户不会困惑。'),
    ('二级标题说明', '每个一级入口下都有文字说明，用户点之前就知道里面有什么。'),
    ('白名单机制', '重跑/清理等控制动作用白名单限定范围，不对所有任务开放。'),
    ('二次确认原则', '真实执行前必须经过确认层，防止误触。'),
    ('脱敏原则', '分享给他人时，所有 API Key、Token、个人路径、Open ID 替换为占位符。'),
]
for name, desc in principles:
    p = doc.add_paragraph()
    p.add_run(f'• {name}：').bold = True
    p.add_run(desc)

# 扩展方法
doc.add_heading('九、如何扩展新专题', level=2)
doc.add_paragraph('1. 新建面板脚本（如 xxx_panel_card.py），生成对应卡片 JSON。')
doc.add_paragraph('2. 在 cockpit_v5_hub.py 的 HUBS 列表里注册路由。')
doc.add_paragraph('3. 在 OpenClaw 的 quick_actions 处理器里注册命令处理逻辑。')
doc.add_paragraph('4. 调试：先用 python 脚本生成 JSON，发送到飞书验证按钮是否可点。')

# 飞书卡片限制
doc.add_heading('十、飞书卡片大小限制处理', level=2)
doc.add_paragraph('飞书单条卡片约 4KB 限制，完整 Hub 约 11KB。')
doc.add_paragraph('解决：卡片内文字用空格代替换行（避免字符膨胀），精简非核心内容后发精简版，专题内容发独立卡片。')

out = '/Users/zhaoruicn/.openclaw/workspace/summary/从修复飞书卡片到主控制台体系_脱敏版.docx'
doc.save(out)
print(f'✅ 已生成：{out}')
