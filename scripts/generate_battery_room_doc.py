#!/usr/bin/env python3
"""生成数据中心锂电UPS产品设计要求推导文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

def set_chinese_font(run, font_name='宋体', font_size=11):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(10)
                run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
    
    return table

def main():
    doc = Document()
    
    # 标题
    title = doc.add_heading('数据中心锂电UPS产品设计要求推导', 0)
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题/日期
    p = doc.add_paragraph()
    run = p.add_run(f'推导日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.name = '宋体'
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ==================== 1. 规范依据 ====================
    add_heading(doc, '一、规范依据', 1)
    
    headers = ['规范/标准', '编号', '主要内容']
    rows = [
        ['数据中心设计规范', 'GB 50174-2017', '数据中心分级、机房设计要求'],
        ['建筑结构荷载规范', 'GB 50009-2012', '楼板活荷载设计值、荷载组合'],
        ['通信用阀控式密封铅酸蓄电池', 'YD/T 799', '电池规格、尺寸、重量'],
        ['蓄电池选用与安装', '14D202-1', '电池室布置、间距要求'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    # ==================== 2. 楼板荷载分析 ====================
    add_heading(doc, '二、楼板荷载分析', 1)
    
    add_heading(doc, '2.1 数据中心楼板活荷载标准', 2)
    
    headers = ['等级', 'GB 50174-2017要求', '行业常规取值', '适用场景']
    rows = [
        ['A级', '≥500 kg/m²', '800-1000 kg/m²', '重要数据中心'],
        ['B级', '≥300 kg/m²', '400-600 kg/m²', '一般数据中心'],
        ['C级', '无强制要求', '200-400 kg/m²', '简易数据中心'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_heading(doc, '2.2 荷载理解与组合', 2)
    add_paragraph(doc, '根据GB 50009-2012《建筑结构荷载规范》：')
    add_paragraph(doc, '• 楼板活荷载设计值（如800 kg/m²）是指楼板能够承受的设备净荷载')
    add_paragraph(doc, '• 楼板自重由结构梁柱承担，不占用设备荷载额度')
    add_paragraph(doc, '• 承载力验算：S = 1.2Gk + 1.4Qk（永久荷载+可变荷载）')
    add_paragraph(doc, '• 设备重量直接使用楼板活荷载额度，无需额外扣除')
    doc.add_paragraph()
    
    # ==================== 3. 机柜尺寸约束 ====================
    add_heading(doc, '三、机柜尺寸约束条件', 1)
    
    add_heading(doc, '3.1 高度约束', 2)
    headers = ['参数', '数值', '说明']
    rows = [
        ['标准42U机柜', '2000mm', '含脚轮、顶盖'],
        ['47U加高机柜', '2200mm', '部分大型UPS采用'],
        ['主流锂电UPS高度', '2000mm', '与标准42U机柜一致'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_heading(doc, '3.2 宽度约束', 2)
    headers = ['参数', '数值', '说明']
    rows = [
        ['标准19英寸机柜', '600mm', '符合EIA-RS-310-D标准'],
        ['常用机柜深度', '1000mm / 1100mm / 1200mm', '根据设备深度选择'],
        ['主流锂电UPS宽度', '600mm', '与标准机柜一致'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_heading(doc, '3.3 深度约束', 2)
    headers = ['约束因素', '要求', '备注']
    rows = [
        ['维护通道', '≥1000mm', 'GB 50174-2017要求'],
        ['电池架间距', '≥800mm', '便于维护操作'],
        ['设备深度', '≤1000-1200mm', '考虑门开启、操作空间'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    # ==================== 4. 重量推导 ====================
    add_heading(doc, '四、重量限制推导', 1)
    
    add_heading(doc, '4.1 计算公式', 2)
    add_paragraph(doc, '单柜最大重量 = 楼板活荷载标准值 × 单柜占地面积')
    add_paragraph(doc, '')
    add_paragraph(doc, '其中：单柜占地面积 = 机柜宽度 × 机柜深度')
    doc.add_paragraph()
    
    add_heading(doc, '4.2 不同楼板荷载下的重量限制', 2)
    headers = ['楼板活荷载(kg/m²)', '单柜占地(m²)', '单柜最大重量(kg)', '适用等级']
    rows = [
        ['600', '0.6×1.0=0.6', '360', 'B级数据中心'],
        ['800', '0.6×1.0=0.6', '480', 'A级数据中心（普通）'],
        ['1000', '0.6×1.0=0.6', '600', 'A级数据中心（高标准）'],
        ['1200', '0.6×1.0=0.6', '720', '特殊加固楼板'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_heading(doc, '4.3 推荐重量目标', 2)
    headers = ['设计目标', '重量限制', '安全余量']
    rows = [
        ['保守设计', '480kg', '满足800kg/m²楼板'],
        ['推荐设计', '600kg', '满足1000kg/m²楼板'],
        ['高标准设计', '720kg', '需确认楼板≥1200kg/m²'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    # ==================== 5. 产品规格汇总 ====================
    add_heading(doc, '五、产品设计规格汇总', 1)
    
    headers = ['参数', '目标值', '允许范围', '备注']
    rows = [
        ['宽度', '600mm', '≤600mm', '与标准机柜一致'],
        ['高度', '2000mm', '≤2000mm', '适配42U标准机柜空间'],
        ['深度', '1000-1200mm', '1000-1200mm', '根据容量选择'],
        ['重量', '≤600kg', '≤600kg', '满足A级数据中心要求'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_heading(doc, '5.1 容量与重量关系估算', 2)
    headers = ['系统容量', '锂电池重量估算', 'UPS主机重量', '总重量估算']
    rows = [
        ['100kVA / 200kWh', '150-200kg', '80-100kg', '250-300kg ✅'],
        ['200kVA / 400kWh', '300-350kg', '120-150kg', '420-500kg ✅'],
        ['300kVA / 600kWh', '450-500kg', '150-180kg', '600-680kg ⚠️'],
        ['500kVA / 1000kWh', '700-800kg', '200-250kg', '900-1050kg ❌'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    add_paragraph(doc, '注：✅ 表示可满足600kg目标；⚠️ 表示需优化；❌ 表示需分柜或分布式设计')
    doc.add_paragraph()
    
    # ==================== 6. 电池室设计建议 ====================
    add_heading(doc, '六、电池室设计建议', 1)
    
    headers = ['项目', '设计要求', '依据']
    rows = [
        ['楼板荷载', '≥1000 kg/m²', '预留余量，便于扩容'],
        ['通道宽度', '≥1000mm', '便于维护操作'],
        ['机柜间距', '≥800mm', '电池架间距要求'],
        ['环境温度', '20-30℃', '延长电池寿命'],
        ['相对湿度', '45%-65%', '防潮要求'],
        ['通风换气', '4-6次/小时', '铅酸电池要求（锂电可选）'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    # ==================== 7. 结论 ====================
    add_heading(doc, '七、结论', 1)
    
    p = doc.add_paragraph()
    run = p.add_run('基于上述推导，数据中心锂电UPS产品的设计规格如下：')
    run.font.name = '宋体'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    headers = ['项目', '目标值', '备注']
    rows = [
        ['宽度', '600mm', '标准19英寸机柜'],
        ['高度', '2000mm', '适配42U机柜空间'],
        ['深度', '1000-1200mm', '根据容量选择'],
        ['重量', '≤600kg', '满足1000kg/m²楼板荷载'],
        ['推荐容量', '≤300kVA/600kWh', '单柜重量控制在600kg以内'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()
    
    add_paragraph(doc, '以上规格可满足大多数A级数据中心的安装要求，对于高密度场景或特殊需求，请与甲方确认具体楼板荷载参数。')
    
    # 保存
    output_path = '/Users/zhaoruicn/.openclaw/workspace/数据中心锂电UPS产品设计要求推导.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')

if __name__ == '__main__':
    main()
